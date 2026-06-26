"""
Exportação de projetos e roteiros para DOCX e PDF, incluindo capa, formatação e estatísticas.

"""

from __future__ import annotations

import html
import io
import re
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Iterable

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle, StyleSheet1, getSampleStyleSheet
from reportlab.lib.units import cm, inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from .manipulador_capitulos import ler_capitulo
from .manipulador_roteiros import ler_roteiro
from .manipulador_projetos import caminho_capa_projeto, obter_pasta_projeto, obter_projeto
from .utilidades import (
    inferir_extensao_data_url,
    nome_seguro,
    obter_configuracoes,
    obter_pasta_exportacao_projeto,
)
from .logging_config import obter_logger, registrar_etapa
logger = obter_logger(__name__)

FONTES_DOCX = {
    "fonte-serif": "Times New Roman",
    "fonte-sans": "Arial",
    "fonte-mono": "Consolas",
}

FONTES_RL = {
    ("fonte-serif", False, False): "Times-Roman",
    ("fonte-serif", True, False): "Times-Bold",
    ("fonte-serif", False, True): "Times-Italic",
    ("fonte-serif", True, True): "Times-BoldItalic",
    ("fonte-sans", False, False): "Helvetica",
    ("fonte-sans", True, False): "Helvetica-Bold",
    ("fonte-sans", False, True): "Helvetica-Oblique",
    ("fonte-sans", True, True): "Helvetica-BoldOblique",
    ("fonte-mono", False, False): "Courier",
    ("fonte-mono", True, False): "Courier-Bold",
    ("fonte-mono", False, True): "Courier-Oblique",
    ("fonte-mono", True, True): "Courier-BoldOblique",
}

# Padrão profissional de roteiro: papel Letter, Courier 12 pt, margens
# 1,5" à esquerda e 1" nas demais bordas. Os recuos abaixo são
# relativos à área útil do texto; somados à margem esquerda, reproduzem
# as posições tradicionais de cena/ação, personagem, rubrica e diálogo.
ROTEIRO_FONTE_DOCX = "Courier New"
ROTEIRO_FONTE_PDF = "Courier"
ROTEIRO_TAMANHO_PT = 12
ROTEIRO_ENTRELINHA_PT = 12
ROTEIRO_ESPACO_PADRAO_PT = 12
ROTEIRO_RECUOS_POL = {
    "scene_heading": (0.0, 0.0),
    "action": (0.0, 0.0),
    "shot": (0.0, 0.0),
    "music": (0.0, 0.0),
    "transition": (0.0, 0.0),
    "character": (2.2, 0.0),
    "dialogue": (1.0, 1.5),
    "parenthetical": (1.5, 1.9),
}
ROTEIRO_ESPACOS_DEPOIS_PT = {
    "scene_heading": 12,
    "action": 12,
    "shot": 12,
    "music": 12,
    "transition": 12,
    "character": 0,
    "dialogue": 12,
    "parenthetical": 0,
}


@registrar_etapa
def _remover_links_recursos_exportacao(soup: BeautifulSoup) -> BeautifulSoup:
    """Remove links internos de recursos, preservando somente o conteúdo textual/inline."""
    for tag in list(soup.find_all("a")):
        classes = tag.get("class") or []
        if isinstance(classes, str):
            classes = classes.split()
        if "editor-recurso-link" in classes or tag.has_attr("data-recurso-tipo") or tag.has_attr("data-recurso-nome"):
            tag.unwrap()
    return soup


@registrar_etapa
def _fragmento_html(html_fragmento: str) -> BeautifulSoup:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        html_fragmento: Valor usado pela rotina para compor a operação de fragmento html.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    soup = BeautifulSoup(f"<div>{html_fragmento or '<p><br></p>'}</div>", "html.parser")
    return _remover_links_recursos_exportacao(soup)


@registrar_etapa
def _iterar_blocos(soup: BeautifulSoup) -> Iterable[Tag]:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        soup: Valor usado pela rotina para compor a operação de iterar blocos.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    raiz = soup.div or soup
    for child in raiz.children:
        if isinstance(child, NavigableString):
            if str(child).strip():
                p = soup.new_tag("p")
                p.string = str(child)
                yield p
            continue
        if not isinstance(child, Tag):
            continue
        if child.name in {"p", "h1", "h2", "h3", "ul", "ol", "blockquote", "img", "hr"}:
            yield child
        elif child.name == "div" and child.get("data-block-type"):
            yield child
        elif child.name == "div":
            # blocos soltos acabam virando parágrafo
            if child.find(["p", "h1", "h2", "h3", "ul", "ol", "blockquote", "img", "hr"], recursive=False):
                for sub in child.find_all(["p", "h1", "h2", "h3", "ul", "ol", "blockquote", "img", "hr"], recursive=False):
                    yield sub
            else:
                p = soup.new_tag("p")
                for sub in list(child.contents):
                    p.append(sub)
                yield p
        else:
            p = soup.new_tag("p")
            p.append(child)
            yield p


@registrar_etapa
def _primeiro_wrapper(bloco: Tag | None) -> Tag | None:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        bloco: Bloco de conteúdo extraído ou gerado durante o processamento.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    if not bloco:
        return None
    for child in bloco.children:
        if isinstance(child, Tag) and child.name == "span" and "editor-formatacao" in (child.get("class") or []):
            return child
    return None


@registrar_etapa
def _classes_bloco(bloco: Tag) -> set[str]:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        bloco: Bloco de conteúdo extraído ou gerado durante o processamento.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    wrapper = _primeiro_wrapper(bloco)
    classes = set(wrapper.get("class") or []) if wrapper else set(bloco.get("class") or [])
    classes.update(set(bloco.get("class") or []))
    classes.discard("editor-formatacao")
    return classes


@registrar_etapa
def _tipo_bloco_roteiro(bloco: Tag) -> str:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        bloco: Bloco de conteúdo extraído ou gerado durante o processamento.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    return (bloco.get("data-block-type") or "").strip().lower()


@registrar_etapa
def _texto_bloco_maiusculo(bloco: Tag) -> bool:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        bloco: Bloco de conteúdo extraído ou gerado durante o processamento.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    return _tipo_bloco_roteiro(bloco) in {"scene_heading", "character", "shot", "transition"}


@registrar_etapa
def _conteudo_inline(bloco: Tag) -> list:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        bloco: Bloco de conteúdo extraído ou gerado durante o processamento.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    wrapper = _primeiro_wrapper(bloco)
    return list((wrapper or bloco).contents)


@registrar_etapa
def _classe_fonte(classes: set[str]) -> str:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        classes: Valor usado pela rotina para compor a operação de classe fonte.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    for nome in ["fonte-sans", "fonte-serif", "fonte-mono"]:
        if nome in classes:
            return nome
    return "fonte-serif"


@registrar_etapa
def _eh_separador_decorativo(classes: set[str]) -> bool:
    """
    Avalia uma condição específica do conteúdo recebido e devolve uma resposta simples para orientar o fluxo.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        classes: Valor usado pela rotina para compor a operação de eh separador decorativo.

    Retorno:
        True ou False, conforme a condição analisada seja atendida.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    return "editor-separador-decorativo" in classes


@registrar_etapa
def _alinhamento_docx(classes: set[str], padrao: str = "justify"):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        classes: Valor usado pela rotina para compor a operação de alinhamento docx.
        padrao: Valor usado pela rotina para compor a operação de alinhamento docx.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    if "alinhar-a-esquerda" in classes:
        return WD_ALIGN_PARAGRAPH.LEFT
    if "alinhar-centro" in classes:
        return WD_ALIGN_PARAGRAPH.CENTER
    if "alinhar-a-direita" in classes:
        return WD_ALIGN_PARAGRAPH.RIGHT
    if "justificar" in classes or padrao == "justify":
        return WD_ALIGN_PARAGRAPH.JUSTIFY
    if padrao == "left":
        return WD_ALIGN_PARAGRAPH.LEFT
    if padrao == "center":
        return WD_ALIGN_PARAGRAPH.CENTER
    return WD_ALIGN_PARAGRAPH.JUSTIFY


@registrar_etapa
def _alinhamento_rl(classes: set[str], padrao: str = "justify"):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        classes: Valor usado pela rotina para compor a operação de alinhamento rl.
        padrao: Valor usado pela rotina para compor a operação de alinhamento rl.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    if "alinhar-a-esquerda" in classes:
        return TA_LEFT
    if "alinhar-centro" in classes:
        return TA_CENTER
    if "alinhar-a-direita" in classes:
        return TA_RIGHT
    if "justificar" in classes or padrao == "justify":
        return TA_JUSTIFY
    if padrao == "left":
        return TA_LEFT
    if padrao == "center":
        return TA_CENTER
    return TA_JUSTIFY


@registrar_etapa
def _largura_imagem(caminho: str | Path | io.BytesIO, largura_max_px: int = 1200) -> tuple[io.BytesIO, int, int]:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        caminho: Caminho de arquivo ou pasta usado como origem ou destino da operação.
        largura_max_px: Valor usado pela rotina para compor a operação de largura imagem.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    if isinstance(caminho, io.BytesIO):
        dados = caminho.getvalue()
    elif isinstance(caminho, Path):
        dados = caminho.read_bytes()
    else:
        dados = Path(caminho).read_bytes()
    imagem = Image.open(io.BytesIO(dados)).convert("RGB")
    largura, altura = imagem.size
    if largura > largura_max_px:
        proporcao = largura_max_px / float(largura)
        imagem = imagem.resize((int(largura * proporcao), int(altura * proporcao)), Image.LANCZOS)
        largura, altura = imagem.size
    saida = io.BytesIO()
    imagem.save(saida, format="PNG")
    saida.seek(0)
    return saida, largura, altura


@registrar_etapa
def _fonte_rl(font_class: str, bold: bool, italic: bool) -> str:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        font_class: Valor usado pela rotina para compor a operação de fonte rl.
        bold: Valor usado pela rotina para compor a operação de fonte rl.
        italic: Valor usado pela rotina para compor a operação de fonte rl.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    return FONTES_RL.get((font_class, bold, italic), "Times-Roman")


@registrar_etapa
def _escape_texto(texto: str) -> str:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        texto: Texto bruto ou parcialmente tratado que será analisado pela função.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    return html.escape(texto).replace("\n", "<br/>")


@registrar_etapa
def _estado_inline(base: dict | None = None, **novos) -> dict:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        base: Valor usado pela rotina para compor a operação de estado inline.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    estado = dict(base or {"bold": False, "italic": False, "underline": False, "font_class": "fonte-serif"})
    estado.update(novos)
    return estado


@registrar_etapa
def _run_set_font(run, nome_fonte: str):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        run: Valor usado pela rotina para compor a operação de run set font.
        nome_fonte: Valor usado pela rotina para compor a operação de run set font.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    run.font.name = nome_fonte
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), nome_fonte)


@registrar_etapa
def _docx_adicionar_inline(paragrafo, no, estado: dict, largura_max_pol: float):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        paragrafo: Valor usado pela rotina para compor a operação de docx adicionar inline.
        no: Valor usado pela rotina para compor a operação de docx adicionar inline.
        estado: Valor usado pela rotina para compor a operação de docx adicionar inline.
        largura_max_pol: Valor usado pela rotina para compor a operação de docx adicionar inline.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    if isinstance(no, NavigableString):
        texto = str(no)
        if not texto:
            return
        run = paragrafo.add_run(texto)
        run.bold = estado["bold"]
        run.italic = estado["italic"]
        run.underline = estado["underline"]
        _run_set_font(run, FONTES_DOCX.get(estado["font_class"], "Times New Roman"))
        return
    if not isinstance(no, Tag):
        return

    nome = no.name.lower()
    if nome == "br":
        paragrafo.add_run().add_break()
        return
    if nome == "img":
        imagem_stream = _imagem_de_tag(no)[0]
        if imagem_stream:
            paragrafo.add_run().add_picture(imagem_stream, width=Inches(largura_max_pol))
        return

    prox = dict(estado)
    classes = set(no.get("class") or [])
    if nome in {"strong", "b"}:
        prox["bold"] = True
    if nome in {"em", "i"}:
        prox["italic"] = True
    if nome == "u":
        prox["underline"] = True
    for fonte in ["fonte-sans", "fonte-serif", "fonte-mono"]:
        if fonte in classes:
            prox["font_class"] = fonte

    for child in no.children:
        _docx_adicionar_inline(paragrafo, child, prox, largura_max_pol)


@registrar_etapa
def _imagem_de_tag(tag: Tag) -> tuple[io.BytesIO | None, int, int]:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        tag: Valor usado pela rotina para compor a operação de imagem de tag.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    src = (tag.get("src") or "").strip()
    if not src:
        return None, 0, 0
    try:
        if src.startswith("data:image/"):
            _ext, dados = inferir_extensao_data_url(src)
            return _largura_imagem(io.BytesIO(dados))
        caminho = Path(src)
        if caminho.exists():
            return _largura_imagem(caminho)
    except Exception:
        return None, 0, 0
    return None, 0, 0


@registrar_etapa
def _docx_configurar_paragrafo(paragrafo, classes: set[str], cfg: dict, *, tipo: str = "p"):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        paragrafo: Valor usado pela rotina para compor a operação de docx configurar paragrafo.
        classes: Valor usado pela rotina para compor a operação de docx configurar paragrafo.
        cfg: Configurações de formatação usadas durante a montagem do documento.
        tipo: Valor usado pela rotina para compor a operação de docx configurar paragrafo.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    fmt = paragrafo.paragraph_format
    if tipo == "h1":
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.first_line_indent = Pt(0)
        fmt.space_before = Pt(18)
        fmt.space_after = Pt(12)
        fmt.line_spacing = 1.2
    elif tipo == "h2":
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = Pt(0)
        fmt.space_before = Pt(16)
        fmt.space_after = Pt(10)
        fmt.line_spacing = 1.2
    elif tipo == "h3":
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = Pt(0)
        fmt.space_before = Pt(14)
        fmt.space_after = Pt(8)
        fmt.line_spacing = 1.2
    else:
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER if _eh_separador_decorativo(classes) else _alinhamento_docx(classes, padrao="justify")
        if _eh_separador_decorativo(classes) or "recuo-a-esquerda" in classes:
            fmt.first_line_indent = Pt(0)
        elif "recuo-a-direita" in classes:
            fmt.first_line_indent = Pt(float(cfg["identacao_paragrafo_maior_em"]) * float(cfg["tamanho_fonte_px"]) * 0.75)
        else:
            fmt.first_line_indent = Pt(float(cfg["identacao_paragrafo_em"]) * float(cfg["tamanho_fonte_px"]) * 0.75)
        fmt.space_after = Pt(float(cfg["espaco_paragrafos_em"]) * 10)
        fmt.line_spacing = float(cfg["espaco_linhas"])
    if tipo == "blockquote":
        fmt.left_indent = Pt(28)
        fmt.right_indent = Pt(18)
        fmt.first_line_indent = Pt(0)


@registrar_etapa
def _docx_processar_bloco(doc: Document, bloco: Tag, cfg: dict, largura_max_pol: float, roteiro_profissional: bool = False):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        doc: Documento em processamento, normalmente vindo de uma biblioteca de exportação ou leitura.
        bloco: Bloco de conteúdo extraído ou gerado durante o processamento.
        cfg: Configurações de formatação usadas durante a montagem do documento.
        largura_max_pol: Valor usado pela rotina para compor a operação de docx processar bloco.
        roteiro_profissional: Valor usado pela rotina para compor a operação de docx processar bloco.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    nome = bloco.name.lower()
    classes = _classes_bloco(bloco)
    tipo_roteiro = _tipo_bloco_roteiro(bloco)
    if tipo_roteiro == "comment":
        return

    if nome in {"ul", "ol"}:
        estilo_lista = "List Bullet" if nome == "ul" else "List Number"
        for li in bloco.find_all("li", recursive=False):
            p = doc.add_paragraph(style=estilo_lista)
            _docx_configurar_paragrafo(p, classes, cfg, tipo="p")
            for child in _conteudo_inline(li):
                _docx_adicionar_inline(p, child, _estado_inline(font_class=_classe_fonte(classes)), largura_max_pol)
        return

    if nome == "img":
        imagem_stream, _w, _h = _imagem_de_tag(bloco)
        if imagem_stream:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(imagem_stream, width=Inches(largura_max_pol))
            p.paragraph_format.space_after = Pt(10)
        return

    if nome == "hr":
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        bordas = OxmlElement("w:pBdr")
        topo = OxmlElement("w:top")
        topo.set(qn("w:val"), "single")
        topo.set(qn("w:sz"), "12")
        topo.set(qn("w:space"), "1")
        topo.set(qn("w:color"), "AAB2BF")
        bordas.append(topo)
        p._p.get_or_add_pPr().append(bordas)
        return

    p = doc.add_paragraph()
    _docx_configurar_paragrafo(p, classes, cfg, tipo="blockquote" if nome == "blockquote" else nome)

    if tipo_roteiro:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt = p.paragraph_format
        fmt.first_line_indent = Pt(0)
        fmt.left_indent = Pt(0)
        fmt.right_indent = Pt(0)
        fmt.space_after = Pt(8)
        fmt.line_spacing = 1.0
        if tipo_roteiro == "character":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.left_indent = Pt(110)
            fmt.right_indent = Pt(110)
        elif tipo_roteiro == "dialogue":
            fmt.left_indent = Pt(72)
            fmt.right_indent = Pt(62)
        elif tipo_roteiro == "parenthetical":
            fmt.left_indent = Pt(92)
            fmt.right_indent = Pt(92)
        elif tipo_roteiro == "transition":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif tipo_roteiro == "comment":
            return

    tamanho = 12 if tipo_roteiro else cfg["tamanho_fonte_px"]
    if nome == "h1":
        tamanho = cfg["h1_px"]
    elif nome == "h2":
        tamanho = cfg["h2_px"]
    elif nome == "h3":
        tamanho = cfg["h3_px"]

    base = _estado_inline(font_class=_classe_fonte(classes))
    conteudo_inline = _conteudo_inline(bloco)
    if roteiro_profissional and tipo_roteiro == "scene_heading":
        texto_limpo = _strip_scene_number_prefix_text("".join(str(child) for child in bloco.strings))
        conteudo_inline = [texto_limpo]
    for child in conteudo_inline:
        _docx_adicionar_inline(p, child, base, largura_max_pol)

    if tipo_roteiro in {"scene_heading", "character", "parenthetical"}:
        fmt.keep_with_next = True
        fmt.widow_control = True

    if _texto_bloco_maiusculo(bloco):
        for run in p.runs:
            run.text = run.text.upper()

    for run in p.runs:
        _run_set_font(run, "Courier New" if tipo_roteiro else FONTES_DOCX.get(base["font_class"], "Times New Roman"))
        if not run.font.size:
            run.font.size = Pt(tamanho)
        if nome in {"h1", "h2", "h3"}:
            run.bold = True
        if tipo_roteiro == "comment":
            run.font.color.rgb = None


@registrar_etapa
def _stylesheet_pdf(cfg: dict) -> StyleSheet1:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        cfg: Configurações de formatação usadas durante a montagem do documento.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    base = getSampleStyleSheet()
    estilos = StyleSheet1()
    estilos.add(ParagraphStyle(
        name="Base",
        parent=base["BodyText"],
        fontName="Times-Roman",
        fontSize=float(cfg["tamanho_fonte_px"]),
        leading=float(cfg["tamanho_fonte_px"]) * float(cfg["espaco_linhas"]),
        spaceAfter=float(cfg["tamanho_fonte_px"]) * float(cfg["espaco_paragrafos_em"]) * 0.45,
        firstLineIndent=float(cfg["tamanho_fonte_px"]) * float(cfg["identacao_paragrafo_em"]),
        alignment=TA_JUSTIFY,
    ))
    estilos.add(ParagraphStyle(name="H1", parent=estilos["Base"], fontName="Times-Bold", fontSize=float(cfg["h1_px"]), leading=float(cfg["h1_px"]) * 1.25, alignment=TA_CENTER, firstLineIndent=0, spaceBefore=18, spaceAfter=12))
    estilos.add(ParagraphStyle(name="H2", parent=estilos["Base"], fontName="Times-Bold", fontSize=float(cfg["h2_px"]), leading=float(cfg["h2_px"]) * 1.2, alignment=TA_LEFT, firstLineIndent=0, spaceBefore=16, spaceAfter=10))
    estilos.add(ParagraphStyle(name="H3", parent=estilos["Base"], fontName="Times-Bold", fontSize=float(cfg["h3_px"]), leading=float(cfg["h3_px"]) * 1.2, alignment=TA_LEFT, firstLineIndent=0, spaceBefore=14, spaceAfter=8))
    estilos.add(ParagraphStyle(name="Blockquote", parent=estilos["Base"], firstLineIndent=0, leftIndent=28, rightIndent=18, textColor="#374151"))
    estilos.add(ParagraphStyle(
        name="ScreenplayBase",
        parent=estilos["Base"],
        fontName=ROTEIRO_FONTE_PDF,
        fontSize=ROTEIRO_TAMANHO_PT,
        leading=ROTEIRO_ENTRELINHA_PT,
        firstLineIndent=0,
        spaceBefore=0,
        spaceAfter=ROTEIRO_ESPACO_PADRAO_PT,
        alignment=TA_LEFT,
        wordWrap="LTR",
    ))
    estilos.add(ParagraphStyle(name="ScreenplayScene", parent=estilos["ScreenplayBase"], alignment=TA_LEFT, leftIndent=0, rightIndent=0, spaceAfter=ROTEIRO_ESPACOS_DEPOIS_PT["scene_heading"]))
    estilos.add(ParagraphStyle(name="ScreenplayAction", parent=estilos["ScreenplayBase"], alignment=TA_LEFT, leftIndent=0, rightIndent=0, spaceAfter=ROTEIRO_ESPACOS_DEPOIS_PT["action"]))
    estilos.add(ParagraphStyle(name="ScreenplayCharacter", parent=estilos["ScreenplayBase"], alignment=TA_LEFT, leftIndent=ROTEIRO_RECUOS_POL["character"][0] * inch, rightIndent=ROTEIRO_RECUOS_POL["character"][1] * inch, spaceAfter=ROTEIRO_ESPACOS_DEPOIS_PT["character"]))
    estilos.add(ParagraphStyle(name="ScreenplayDialogue", parent=estilos["ScreenplayBase"], alignment=TA_LEFT, leftIndent=ROTEIRO_RECUOS_POL["dialogue"][0] * inch, rightIndent=ROTEIRO_RECUOS_POL["dialogue"][1] * inch, spaceAfter=ROTEIRO_ESPACOS_DEPOIS_PT["dialogue"]))
    estilos.add(ParagraphStyle(name="ScreenplayParenthetical", parent=estilos["ScreenplayBase"], alignment=TA_LEFT, leftIndent=ROTEIRO_RECUOS_POL["parenthetical"][0] * inch, rightIndent=ROTEIRO_RECUOS_POL["parenthetical"][1] * inch, spaceAfter=ROTEIRO_ESPACOS_DEPOIS_PT["parenthetical"]))
    estilos.add(ParagraphStyle(name="ScreenplayShot", parent=estilos["ScreenplayBase"], alignment=TA_LEFT, leftIndent=0, rightIndent=0, spaceAfter=ROTEIRO_ESPACOS_DEPOIS_PT["shot"]))
    estilos.add(ParagraphStyle(name="ScreenplayTransition", parent=estilos["ScreenplayBase"], alignment=TA_RIGHT, leftIndent=0, rightIndent=0, spaceAfter=ROTEIRO_ESPACOS_DEPOIS_PT["transition"]))
    estilos.add(ParagraphStyle(name="ScreenplayMusic", parent=estilos["ScreenplayBase"], alignment=TA_LEFT, leftIndent=0, rightIndent=0, spaceAfter=ROTEIRO_ESPACOS_DEPOIS_PT["music"]))
    return estilos


@registrar_etapa
def _inline_para_pdf(no, estado: dict) -> str:
    """
    Converte a estrutura recebida para uma representação adequada ao próximo estágio do processamento.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        no: Valor usado pela rotina para compor a operação de inline para pdf.
        estado: Valor usado pela rotina para compor a operação de inline para pdf.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    if isinstance(no, NavigableString):
        return _escape_texto(str(no))
    if not isinstance(no, Tag):
        return ""
    nome = no.name.lower()
    if nome == "br":
        return "<br/>"
    if nome == "img":
        return ""

    prox = dict(estado)
    classes = set(no.get("class") or [])
    if nome in {"strong", "b"}:
        prox["bold"] = True
    if nome in {"em", "i"}:
        prox["italic"] = True
    if nome == "u":
        prox["underline"] = True
    for fonte in ["fonte-sans", "fonte-serif", "fonte-mono"]:
        if fonte in classes:
            prox["font_class"] = fonte

    conteudo = "".join(_inline_para_pdf(child, prox) for child in no.children)
    font_name = _fonte_rl(prox["font_class"], prox["bold"], prox["italic"])
    texto = f'<font name="{font_name}">{conteudo}</font>'
    if prox["underline"]:
        texto = f"<u>{texto}</u>"
    return texto


@registrar_etapa
def _partes_inline_pdf(nos, estado: dict):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        nos: Valor usado pela rotina para compor a operação de partes inline pdf.
        estado: Valor usado pela rotina para compor a operação de partes inline pdf.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    buffer = []
    partes = []

    @registrar_etapa
    def flush_buffer():
        """
        Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.
    
        A função isola esta responsabilidade para deixar o fluxo principal mais fácil
        de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
        normalizações ou verificações necessárias e entrega um resultado previsível
        para a próxima camada do sistema.
    
        Retorno:
            O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.
    
        Observações:
            Mantém a lógica centralizada e evita que detalhes de implementação vazem
            para quem apenas precisa acionar este comportamento.
        """
        texto = ''.join(buffer).strip()
        if texto:
            partes.append(("texto", texto))
        buffer.clear()

    for no in nos:
        if isinstance(no, Tag) and no.name and no.name.lower() == "img":
            flush_buffer()
            partes.append(("imagem", no))
            continue
        buffer.append(_inline_para_pdf(no, estado))

    flush_buffer()
    return partes


@registrar_etapa
def _append_pdf_paragraph_or_image(story: list, estilo, tipo: str, conteudo, largura_imagem_pt: float):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        story: Valor usado pela rotina para compor a operação de append pdf paragraph or image.
        estilo: Valor usado pela rotina para compor a operação de append pdf paragraph or image.
        tipo: Valor usado pela rotina para compor a operação de append pdf paragraph or image.
        conteudo: Valor usado pela rotina para compor a operação de append pdf paragraph or image.
        largura_imagem_pt: Valor usado pela rotina para compor a operação de append pdf paragraph or image.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    if tipo == "imagem":
        imagem_stream, largura, altura = _imagem_de_tag(conteudo)
        if imagem_stream:
            proporcao = altura / float(largura or 1)
            img = RLImage(imagem_stream, width=largura_imagem_pt, height=largura_imagem_pt * proporcao)
            story.append(img)
            story.append(Spacer(1, 8))
        return

    texto = (conteudo or '').strip()
    if not texto:
        texto = "&#160;"
    story.append(Paragraph(texto, estilo))


@registrar_etapa
def _pdf_story_para_capitulo(story: list, titulo: str, html_fragmento: str, cfg: dict, largura_imagem_pt: float, roteiro_profissional: bool = False):
    """
    Converte a estrutura recebida para uma representação adequada ao próximo estágio do processamento.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        story: Valor usado pela rotina para compor a operação de pdf story para capitulo.
        titulo: Título exibido ou salvo para representar o conteúdo tratado.
        html_fragmento: Valor usado pela rotina para compor a operação de pdf story para capitulo.
        cfg: Configurações de formatação usadas durante a montagem do documento.
        largura_imagem_pt: Valor usado pela rotina para compor a operação de pdf story para capitulo.
        roteiro_profissional: Valor usado pela rotina para compor a operação de pdf story para capitulo.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    estilos = _stylesheet_pdf(cfg)
    if titulo:
        story.append(Paragraph(html.escape(titulo), estilos["H1"]))
    soup = _fragmento_html(html_fragmento)
    for bloco in _iterar_blocos(soup):
        nome = bloco.name.lower()
        classes = _classes_bloco(bloco)
        tipo_roteiro = _tipo_bloco_roteiro(bloco)
        if tipo_roteiro == "comment":
            continue
        if nome in {"ul", "ol"}:
            for idx, li in enumerate(bloco.find_all("li", recursive=False), start=1):
                estilo = ParagraphStyle(name=f"ListItem{idx}", parent=estilos["Base"])
                estilo.firstLineIndent = 0
                estilo.leftIndent = 18
                estilo.alignment = _alinhamento_rl(classes, padrao="left")
                texto = "".join(_inline_para_pdf(c, _estado_inline(font_class=_classe_fonte(classes))) for c in _conteudo_inline(li))
                bullet = "•" if nome == "ul" else f"{idx}."
                story.append(Paragraph(texto, estilo, bulletText=bullet))
            continue
        if nome == "img":
            _append_pdf_paragraph_or_image(story, None, "imagem", bloco, largura_imagem_pt)
            continue
        if nome == "hr":
            story.append(Table([[""]], colWidths=[largura_imagem_pt], rowHeights=[2], style=TableStyle([('LINEABOVE', (0,0), (-1,-1), 1.2, colors.HexColor('#AAB2BF'))])))
            story.append(Spacer(1, 8))
            continue

        estilo_nome = "Base"
        if tipo_roteiro:
            mapa_roteiro = {"scene_heading":"ScreenplayScene","action":"ScreenplayAction","character":"ScreenplayCharacter","dialogue":"ScreenplayDialogue","parenthetical":"ScreenplayParenthetical","shot":"ScreenplayShot","transition":"ScreenplayTransition","music":"ScreenplayMusic"}
            estilo_nome = mapa_roteiro.get(tipo_roteiro, "ScreenplayAction")
        if nome == "h1":
            estilo_nome = "H1"
        elif nome == "h2":
            estilo_nome = "H2"
        elif nome == "h3":
            estilo_nome = "H3"
        elif nome == "blockquote":
            estilo_nome = "Blockquote"

        estilo = ParagraphStyle(name=f"{estilo_nome}_{len(story)}", parent=estilos[estilo_nome])
        if nome not in {"h1", "h2", "h3", "blockquote"}:
            estilo.alignment = TA_CENTER if _eh_separador_decorativo(classes) else _alinhamento_rl(classes, padrao="justify")
            if _eh_separador_decorativo(classes) or "recuo-a-esquerda" in classes:
                estilo.firstLineIndent = 0
            elif "recuo-a-direita" in classes:
                estilo.firstLineIndent = float(cfg["tamanho_fonte_px"]) * float(cfg["identacao_paragrafo_maior_em"])
            else:
                estilo.firstLineIndent = float(cfg["tamanho_fonte_px"]) * float(cfg["identacao_paragrafo_em"])

        conteudo_inline = _conteudo_inline(bloco)
        if roteiro_profissional and tipo_roteiro == "scene_heading":
            conteudo_inline = [_strip_scene_number_prefix_text("".join(str(child) for child in bloco.strings))]
        partes = _partes_inline_pdf(conteudo_inline, _estado_inline(font_class=_classe_fonte(classes)))
        if not partes:
            story.append(Paragraph("&#160;", estilo))
            continue

        for tipo, conteudo in partes:
            if tipo == "texto" and _texto_bloco_maiusculo(bloco):
                conteudo = conteudo.upper()
            _append_pdf_paragraph_or_image(story, estilo, tipo, conteudo, largura_imagem_pt)


@registrar_etapa
def _adicionar_capa_docx(doc: Document, projeto: dict):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        doc: Documento em processamento, normalmente vindo de uma biblioteca de exportação ou leitura.
        projeto: Dicionário com os metadados e configurações do projeto atual.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    capa = caminho_capa_projeto(projeto["slug"])
    if capa.exists():
        sec = doc.sections[0]
        largura = sec.page_width - sec.left_margin - sec.right_margin
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(capa), width=largura)
        doc.add_paragraph()


@registrar_etapa
def exportar_projeto_docx(slug: str) -> Path:
    """
    Conduz a exportação dos dados do projeto para o formato esperado, preparando o conteúdo e delegando etapas auxiliares quando necessário.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        slug: Identificador amigável do projeto, livro ou recurso que será manipulado.

    Retorno:
        O caminho do arquivo exportado ou uma estrutura com informações da exportação.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    projeto = obter_projeto(slug)
    pasta_projeto = obter_pasta_projeto(slug)
    pasta_destino = obter_pasta_exportacao_projeto(projeto["titulo"])
    caminho_saida = pasta_destino / f"{nome_seguro(projeto['titulo'])}_manuscrito.docx"

    total_palavras = _contar_palavras_projeto(projeto, pasta_projeto)
    doc = Document()
    sec = doc.sections[0]
    _configurar_secao_manuscrito_docx(sec)
    sec.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    _adicionar_pagina_rosto_manuscrito_docx(doc, projeto, total_palavras)
    sec_body = doc.add_section(WD_SECTION.NEW_PAGE)
    _configurar_secao_manuscrito_docx(sec_body)
    _docx_reiniciar_numeracao_secao(sec_body, 1)
    _configurar_header_manuscrito_docx(sec_body, projeto)
    _adicionar_body_manuscrito_docx(doc, projeto, pasta_projeto)

    doc.save(caminho_saida)
    return caminho_saida


@registrar_etapa
def exportar_projeto_pdf(slug: str) -> Path:
    """
    Conduz a exportação dos dados do projeto para o formato esperado, preparando o conteúdo e delegando etapas auxiliares quando necessário.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        slug: Identificador amigável do projeto, livro ou recurso que será manipulado.

    Retorno:
        O caminho do arquivo exportado ou uma estrutura com informações da exportação.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    projeto = obter_projeto(slug)
    pasta_projeto = obter_pasta_projeto(slug)
    pasta_destino = obter_pasta_exportacao_projeto(projeto["titulo"])
    caminho_saida = pasta_destino / f"{nome_seguro(projeto['titulo'])}_manuscrito.pdf"

    total_palavras = _contar_palavras_projeto(projeto, pasta_projeto)
    doc = SimpleDocTemplate(
        str(caminho_saida),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    estilos = _estilos_manuscrito_pdf()
    story = []
    _adicionar_pagina_rosto_manuscrito_pdf(story, projeto, total_palavras, estilos)
    _adicionar_body_manuscrito_pdf(story, projeto, pasta_projeto, estilos)
    doc.build(
        story,
        onFirstPage=lambda c, d: _pdf_header_manuscrito(c, d, projeto),
        onLaterPages=lambda c, d: _pdf_header_manuscrito(c, d, projeto),
    )
    return caminho_saida


# ===== Exportação técnica de livros em formato de manuscrito =====

@registrar_etapa
def _texto_normalizado_manuscrito(texto: str) -> str:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        texto: Texto bruto ou parcialmente tratado que será analisado pela função.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    return re.sub(r"\s+", " ", str(texto or "")).strip()


@registrar_etapa
def _texto_bloco_manuscrito(bloco: Tag) -> str:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        bloco: Bloco de conteúdo extraído ou gerado durante o processamento.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    if bloco.name == "hr" or _eh_separador_decorativo(_classes_bloco(bloco)):
        return "***"
    if bloco.name == "img":
        return ""
    return _texto_normalizado_manuscrito(bloco.get_text(" ", strip=True))


@registrar_etapa
def _contar_palavras_projeto(projeto: dict, pasta_projeto: Path) -> int:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        projeto: Dicionário com os metadados e configurações do projeto atual.
        pasta_projeto: Pasta raiz do projeto em que os dados relacionados serão lidos ou gravados.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    total = 0
    for cap_meta in projeto.get("capitulos", []):
        try:
            cap = ler_capitulo(pasta_projeto, cap_meta["id"])
        except Exception:
            continue
        html_fragmento = (cap.get("paragrafos") or ["<p><br></p>"])[0]
        soup = _fragmento_html(html_fragmento)
        for bloco in _iterar_blocos(soup):
            total += len(re.findall(r"\b[\wÀ-ÿ'-]+\b", _texto_bloco_manuscrito(bloco), flags=re.UNICODE))
    return total


@registrar_etapa
def _docx_reiniciar_numeracao_secao(sec, inicio: int = 1) -> None:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        sec: Valor usado pela rotina para compor a operação de docx reiniciar numeracao secao.
        inicio: Valor usado pela rotina para compor a operação de docx reiniciar numeracao secao.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    sect_pr = sec._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(max(1, int(inicio or 1))))


@registrar_etapa
def _configurar_secao_manuscrito_docx(sec) -> None:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        sec: Valor usado pela rotina para compor a operação de configurar secao manuscrito docx.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)


@registrar_etapa
def _configurar_header_manuscrito_docx(sec, projeto: dict) -> None:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        sec: Valor usado pela rotina para compor a operação de configurar header manuscrito docx.
        projeto: Dicionário com os metadados e configurações do projeto atual.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    header = sec.header
    try:
        header.is_linked_to_previous = False
    except Exception:
        pass
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    autor = _texto_normalizado_manuscrito(projeto.get("autor") or "")
    titulo = _texto_normalizado_manuscrito(projeto.get("titulo") or "")
    prefixo = " / ".join([v for v in [autor, titulo] if v]) or "Manuscrito"
    run = p.add_run(f"{prefixo} / ")
    _run_set_font(run, "Times New Roman")
    run.font.size = Pt(12)
    page_run = _append_word_field(p, "PAGE")
    _run_set_font(page_run, "Times New Roman")
    page_run.font.size = Pt(12)


@registrar_etapa
def _adicionar_linha_docx(doc: Document, texto: str = "", alinhamento=WD_ALIGN_PARAGRAPH.LEFT, tamanho: float = 12, negrito: bool = False, italico: bool = False, space_after: float = 0):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        doc: Documento em processamento, normalmente vindo de uma biblioteca de exportação ou leitura.
        texto: Texto bruto ou parcialmente tratado que será analisado pela função.
        alinhamento: Valor usado pela rotina para compor a operação de adicionar linha docx.
        tamanho: Valor usado pela rotina para compor a operação de adicionar linha docx.
        negrito: Valor usado pela rotina para compor a operação de adicionar linha docx.
        italico: Valor usado pela rotina para compor a operação de adicionar linha docx.
        space_after: Valor usado pela rotina para compor a operação de adicionar linha docx.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    p = doc.add_paragraph()
    p.alignment = alinhamento
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(texto)
    run.bold = negrito
    run.italic = italico
    run.font.size = Pt(tamanho)
    _run_set_font(run, "Times New Roman")
    return p


@registrar_etapa
def _texto_contagem_palavras_manuscrito(total_palavras: int) -> str:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        total_palavras: Valor usado pela rotina para compor a operação de texto contagem palavras manuscrito.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    return f"aprox. {int(total_palavras):,}".replace(",", ".") + " palavras"


@registrar_etapa
def _adicionar_pagina_rosto_manuscrito_docx(doc: Document, projeto: dict, total_palavras: int) -> None:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        doc: Documento em processamento, normalmente vindo de uma biblioteca de exportação ou leitura.
        projeto: Dicionário com os metadados e configurações do projeto atual.
        total_palavras: Valor usado pela rotina para compor a operação de adicionar pagina rosto manuscrito docx.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    contato = _texto_normalizado_manuscrito(projeto.get("contatos") or "")
    if contato:
        _adicionar_linha_docx(doc, contato, WD_ALIGN_PARAGRAPH.LEFT)
    else:
        _adicionar_linha_docx(doc, "", WD_ALIGN_PARAGRAPH.LEFT)
    if total_palavras:
        _adicionar_linha_docx(doc, _texto_contagem_palavras_manuscrito(total_palavras), WD_ALIGN_PARAGRAPH.RIGHT)

    for _ in range(8):
        doc.add_paragraph()

    titulo = _texto_normalizado_manuscrito(projeto.get("titulo") or "Sem título").upper()
    _adicionar_linha_docx(doc, titulo, WD_ALIGN_PARAGRAPH.CENTER, negrito=True, space_after=6)
    autor = _texto_normalizado_manuscrito(projeto.get("autor") or "")
    if autor:
        _adicionar_linha_docx(doc, f"por {autor}", WD_ALIGN_PARAGRAPH.CENTER)

@registrar_etapa
def _docx_adicionar_paragrafo_manuscrito(doc: Document, texto: str, tipo: str = "p") -> None:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        doc: Documento em processamento, normalmente vindo de uma biblioteca de exportação ou leitura.
        texto: Texto bruto ou parcialmente tratado que será analisado pela função.
        tipo: Valor usado pela rotina para compor a operação de docx adicionar paragrafo manuscrito.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    if not texto:
        return
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.line_spacing = 2.0
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    run = p.add_run(texto)
    _run_set_font(run, "Times New Roman")
    run.font.size = Pt(12)
    if tipo in {"h1", "h2", "h3"} or texto == "***":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.first_line_indent = Inches(0)
        run.bold = tipo in {"h1", "h2"}
    elif tipo == "blockquote":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.left_indent = Inches(0.5)
        fmt.right_indent = Inches(0.5)
        fmt.first_line_indent = Inches(0)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = Inches(0.5)


@registrar_etapa
def _adicionar_body_manuscrito_docx(doc: Document, projeto: dict, pasta_projeto: Path) -> None:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        doc: Documento em processamento, normalmente vindo de uma biblioteca de exportação ou leitura.
        projeto: Dicionário com os metadados e configurações do projeto atual.
        pasta_projeto: Pasta raiz do projeto em que os dados relacionados serão lidos ou gravados.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    for indice, cap_meta in enumerate(projeto.get("capitulos", []), start=1):
        if indice > 1:
            doc.add_page_break()
        cap = ler_capitulo(pasta_projeto, cap_meta["id"])
        titulo_cap = _texto_normalizado_manuscrito(cap.get("titulo") or f"Capítulo {indice}")
        _docx_adicionar_paragrafo_manuscrito(doc, titulo_cap.upper(), "h1")
        doc.add_paragraph()
        html_fragmento = (cap.get("paragrafos") or ["<p><br></p>"])[0]
        soup = _fragmento_html(html_fragmento)
        for bloco in _iterar_blocos(soup):
            texto = _texto_bloco_manuscrito(bloco)
            tipo = bloco.name if bloco.name in {"h1", "h2", "h3", "blockquote"} else "p"
            _docx_adicionar_paragrafo_manuscrito(doc, texto, tipo)


@registrar_etapa
def _pdf_header_manuscrito(canvas, doc, projeto: dict) -> None:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        canvas: Valor usado pela rotina para compor a operação de pdf header manuscrito.
        doc: Documento em processamento, normalmente vindo de uma biblioteca de exportação ou leitura.
        projeto: Dicionário com os metadados e configurações do projeto atual.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    if canvas.getPageNumber() <= 1:
        return
    canvas.saveState()
    canvas.setFont("Times-Roman", 12)
    autor = _texto_normalizado_manuscrito(projeto.get("autor") or "")
    titulo = _texto_normalizado_manuscrito(projeto.get("titulo") or "")
    prefixo = " / ".join([v for v in [autor, titulo] if v]) or "Manuscrito"
    canvas.drawRightString(letter[0] - inch, letter[1] - 0.55 * inch, f"{prefixo} / {canvas.getPageNumber() - 1}")
    canvas.restoreState()


@registrar_etapa
def _estilos_manuscrito_pdf() -> StyleSheet1:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    estilos = StyleSheet1()
    estilos.add(ParagraphStyle("Base", fontName="Times-Roman", fontSize=12, leading=24, alignment=TA_LEFT, firstLineIndent=0.5 * inch, spaceBefore=0, spaceAfter=0))
    estilos.add(ParagraphStyle("Contato", fontName="Times-Roman", fontSize=12, leading=14, alignment=TA_LEFT, firstLineIndent=0, spaceBefore=0, spaceAfter=0))
    estilos.add(ParagraphStyle("Contagem", fontName="Times-Roman", fontSize=12, leading=14, alignment=TA_RIGHT, firstLineIndent=0, spaceBefore=0, spaceAfter=0))
    estilos.add(ParagraphStyle("Titulo", fontName="Times-Bold", fontSize=12, leading=24, alignment=TA_CENTER, firstLineIndent=0, spaceBefore=0, spaceAfter=0))
    estilos.add(ParagraphStyle("Centro", fontName="Times-Roman", fontSize=12, leading=24, alignment=TA_CENTER, firstLineIndent=0, spaceBefore=0, spaceAfter=0))
    estilos.add(ParagraphStyle("Meta", fontName="Times-Roman", fontSize=12, leading=18, alignment=TA_LEFT, firstLineIndent=0, leftIndent=0.5 * inch, rightIndent=0.5 * inch, spaceBefore=0, spaceAfter=6))
    estilos.add(ParagraphStyle("Blockquote", fontName="Times-Roman", fontSize=12, leading=24, alignment=TA_LEFT, firstLineIndent=0, leftIndent=0.5 * inch, rightIndent=0.5 * inch, spaceBefore=0, spaceAfter=0))
    return estilos


@registrar_etapa
def _adicionar_pagina_rosto_manuscrito_pdf(story: list, projeto: dict, total_palavras: int, estilos: StyleSheet1) -> None:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        story: Valor usado pela rotina para compor a operação de adicionar pagina rosto manuscrito pdf.
        projeto: Dicionário com os metadados e configurações do projeto atual.
        total_palavras: Valor usado pela rotina para compor a operação de adicionar pagina rosto manuscrito pdf.
        estilos: Valor usado pela rotina para compor a operação de adicionar pagina rosto manuscrito pdf.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    contato = _texto_normalizado_manuscrito(projeto.get("contatos") or "")
    contagem = _texto_contagem_palavras_manuscrito(total_palavras) if total_palavras else ""
    topo = Table(
        [[Paragraph(html.escape(contato), estilos["Contato"]), Paragraph(html.escape(contagem), estilos["Contagem"])]],
        colWidths=[3.0 * inch, 3.5 * inch],
        hAlign="LEFT",
    )
    topo.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0), ("TOPPADDING", (0, 0), (-1, -1), 0), ("BOTTOMPADDING", (0, 0), (-1, -1), 0)]))
    story.append(topo)
    story.append(Spacer(1, 3.0 * inch))
    story.append(Paragraph(html.escape(_texto_normalizado_manuscrito(projeto.get("titulo") or "Sem título").upper()), estilos["Titulo"]))
    autor = _texto_normalizado_manuscrito(projeto.get("autor") or "")
    if autor:
        story.append(Paragraph(html.escape(f"por {autor}"), estilos["Centro"]))
    story.append(PageBreak())

@registrar_etapa
def _adicionar_body_manuscrito_pdf(story: list, projeto: dict, pasta_projeto: Path, estilos: StyleSheet1) -> None:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        story: Valor usado pela rotina para compor a operação de adicionar body manuscrito pdf.
        projeto: Dicionário com os metadados e configurações do projeto atual.
        pasta_projeto: Pasta raiz do projeto em que os dados relacionados serão lidos ou gravados.
        estilos: Valor usado pela rotina para compor a operação de adicionar body manuscrito pdf.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    for indice, cap_meta in enumerate(projeto.get("capitulos", []), start=1):
        if indice > 1:
            story.append(PageBreak())
        cap = ler_capitulo(pasta_projeto, cap_meta["id"])
        titulo_cap = _texto_normalizado_manuscrito(cap.get("titulo") or f"Capítulo {indice}").upper()
        story.append(Paragraph(html.escape(titulo_cap), estilos["Titulo"]))
        story.append(Spacer(1, 24))
        html_fragmento = (cap.get("paragrafos") or ["<p><br></p>"])[0]
        soup = _fragmento_html(html_fragmento)
        for bloco in _iterar_blocos(soup):
            texto = _texto_bloco_manuscrito(bloco)
            if not texto:
                continue
            if bloco.name in {"h1", "h2", "h3"} or texto == "***":
                estilo = estilos["Titulo"]
            elif bloco.name == "blockquote":
                estilo = estilos["Blockquote"]
            else:
                estilo = estilos["Base"]
            story.append(Paragraph(html.escape(texto), estilo))


@registrar_etapa
def _strip_scene_number_prefix_text(texto: str) -> str:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        texto: Texto bruto ou parcialmente tratado que será analisado pela função.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    return __import__("re").sub(r"^\s*\d+\.\s*", "", str(texto or "")).strip()


@registrar_etapa
def _append_word_field(paragraph, field_name: str):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        paragraph: Valor usado pela rotina para compor a operação de append word field.
        field_name: Valor usado pela rotina para compor a operação de append word field.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    text_el = OxmlElement("w:t")
    text_el.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.extend([fld_begin, instr, fld_sep, text_el, fld_end])
    return run


@registrar_etapa
def _configurar_tab_direita(paragraph, posicao):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        paragraph: Valor usado pela rotina para compor a operação de configurar tab direita.
        posicao: Valor usado pela rotina para compor a operação de configurar tab direita.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    ppr = paragraph._element.get_or_add_pPr()
    tabs = ppr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        ppr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(posicao)))
    tabs.append(tab)


@registrar_etapa
def _configurar_secao_roteiro_docx(sec, cabecalho: str = "", rodape: str = ""):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        sec: Valor usado pela rotina para compor a operação de configurar secao roteiro docx.
        cabecalho: Valor usado pela rotina para compor a operação de configurar secao roteiro docx.
        rodape: Valor usado pela rotina para compor a operação de configurar secao roteiro docx.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.5)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Inches(0.5)
    sec.footer_distance = Inches(0.5)
    sec.different_first_page_header_footer = True

    right_pos = sec.page_width - sec.right_margin

    first_header = sec.first_page_header.paragraphs[0]
    first_header.clear()
    first_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    first_header.paragraph_format.space_after = Pt(0)

    first_footer = sec.first_page_footer.paragraphs[0]
    first_footer.clear()
    first_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first_footer.paragraph_format.space_after = Pt(0)

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    _configurar_tab_direita(header, right_pos)
    if cabecalho:
        run = header.add_run(cabecalho.strip().upper())
        run.font.size = Pt(10)
        _run_set_font(run, "Courier New")
    run_tab = header.add_run("	")
    _run_set_font(run_tab, "Courier New")
    run_page = _append_word_field(header, "PAGE")
    run_page.font.size = Pt(12)
    _run_set_font(run_page, "Courier New")
    run_dot = header.add_run(".")
    run_dot.font.size = Pt(12)
    _run_set_font(run_dot, "Courier New")

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    if rodape:
        run = footer.add_run(rodape.strip().upper())
        run.font.size = Pt(10)
        _run_set_font(run, "Courier New")


@registrar_etapa
def _desenhar_cabecalho_rodape_roteiro_pdf(canvas, doc, cabecalho: str = "", rodape: str = "", mostrar_numero: bool = True, deslocamento_numero: int = 0):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        canvas: Valor usado pela rotina para compor a operação de desenhar cabecalho rodape roteiro pdf.
        doc: Documento em processamento, normalmente vindo de uma biblioteca de exportação ou leitura.
        cabecalho: Valor usado pela rotina para compor a operação de desenhar cabecalho rodape roteiro pdf.
        rodape: Valor usado pela rotina para compor a operação de desenhar cabecalho rodape roteiro pdf.
        mostrar_numero: Valor usado pela rotina para compor a operação de desenhar cabecalho rodape roteiro pdf.
        deslocamento_numero: Valor usado pela rotina para compor a operação de desenhar cabecalho rodape roteiro pdf.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    canvas.saveState()
    canvas.setFont("Courier", 10)
    largura, altura = doc.pagesize
    topo_y = altura - 0.5 * inch
    rodape_y = 0.5 * inch
    if cabecalho:
        canvas.drawString(doc.leftMargin, topo_y, cabecalho.strip().upper())
    if mostrar_numero:
        canvas.setFont("Courier", 12)
        numero = max(1, canvas.getPageNumber() - int(deslocamento_numero or 0))
        canvas.drawRightString(largura - doc.rightMargin, topo_y, f"{numero}.")
        canvas.setFont("Courier", 10)
    if rodape:
        canvas.drawCentredString(largura / 2.0, rodape_y, rodape.strip().upper())
    canvas.restoreState()


@registrar_etapa
def _criar_callbacks_roteiro_pdf(cabecalho: str = "", rodape: str = "", ocultar_primeira: bool = False, deslocamento_numero: int = 0):
    """
    Cria um novo registro, arquivo ou estrutura interna mantendo o padrão de armazenamento do projeto.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        cabecalho: Valor usado pela rotina para compor a operação de criar callbacks roteiro pdf.
        rodape: Valor usado pela rotina para compor a operação de criar callbacks roteiro pdf.
        ocultar_primeira: Valor usado pela rotina para compor a operação de criar callbacks roteiro pdf.
        deslocamento_numero: Valor usado pela rotina para compor a operação de criar callbacks roteiro pdf.

    Retorno:
        Os dados do novo item criado, incluindo identificadores gerados quando existirem.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    @registrar_etapa
    def primeira_pagina(canvas, doc):
        """
        Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.
    
        A função isola esta responsabilidade para deixar o fluxo principal mais fácil
        de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
        normalizações ou verificações necessárias e entrega um resultado previsível
        para a próxima camada do sistema.
    
        Parâmetros:
            canvas: Valor usado pela rotina para compor a operação de primeira pagina.
            doc: Documento em processamento, normalmente vindo de uma biblioteca de exportação ou leitura.
    
        Retorno:
            O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.
    
        Observações:
            Mantém a lógica centralizada e evita que detalhes de implementação vazem
            para quem apenas precisa acionar este comportamento.
        """
        if ocultar_primeira:
            return
        _desenhar_cabecalho_rodape_roteiro_pdf(canvas, doc, cabecalho, rodape, mostrar_numero=False, deslocamento_numero=deslocamento_numero)

    @registrar_etapa
    def demais_paginas(canvas, doc):
        """
        Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.
    
        A função isola esta responsabilidade para deixar o fluxo principal mais fácil
        de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
        normalizações ou verificações necessárias e entrega um resultado previsível
        para a próxima camada do sistema.
    
        Parâmetros:
            canvas: Valor usado pela rotina para compor a operação de demais paginas.
            doc: Documento em processamento, normalmente vindo de uma biblioteca de exportação ou leitura.
    
        Retorno:
            O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.
    
        Observações:
            Mantém a lógica centralizada e evita que detalhes de implementação vazem
            para quem apenas precisa acionar este comportamento.
        """
        _desenhar_cabecalho_rodape_roteiro_pdf(canvas, doc, cabecalho, rodape, mostrar_numero=True, deslocamento_numero=deslocamento_numero)

    return primeira_pagina, demais_paginas


@registrar_etapa
def _coletar_dados_roteiro_profissional(projeto: dict, roteiro: dict) -> dict:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        projeto: Dicionário com os metadados e configurações do projeto atual.
        roteiro: Valor usado pela rotina para compor a operação de coletar dados roteiro profissional.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    return {
        "titulo_projeto": (projeto.get("titulo") or "").strip(),
        "autor": (projeto.get("autor") or "").strip(),
        "contatos": (projeto.get("contatos") or "").strip(),
        "informacoes_adicionais": (projeto.get("informacoes_adicionais") or "").strip(),
        "titulo_roteiro": (roteiro.get("titulo") or "").strip(),
        "cabecalho": (roteiro.get("cabecalho") or "").strip(),
        "rodape": (roteiro.get("rodape") or "").strip(),
        "html": (roteiro.get("paragrafos") or ["<p><br></p>"])[0],
        "logline": (roteiro.get("logline") or "").strip(),
        "sinopse": (roteiro.get("sinopse") or "").strip(),
        "genero": (roteiro.get("genero") or "").strip(),
        "tipo_roteiro": (roteiro.get("tipo_roteiro") or "spec_script").strip(),
        "tipo_roteiro_label": (roteiro.get("tipo_roteiro_label") or t("script.spec_script")).strip(),
        "copyright": bool(roteiro.get("copyright", (roteiro.get("tipo_roteiro") or "spec_script") != "spec_script")),
        "prefixo_cena": str(roteiro.get("prefixo_cena") or "0").strip() or "0",
        "numeracao_inicial": max(1, int(roteiro.get("numeracao_inicial") or 1)),
        "catalogo_personagens": list(roteiro.get("catalogo_personagens") or []),
        "catalogo_locais": list(roteiro.get("catalogo_locais") or []),
    }


@registrar_etapa
def _tipo_exportacao_roteiro(dados: dict) -> dict:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        dados: Valor usado pela rotina para compor a operação de tipo exportacao roteiro.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    tipo = str(dados.get("tipo_roteiro") or "spec_script").strip().lower()
    label = str(dados.get("tipo_roteiro_label") or t("script.spec_script")).strip()
    scene_numbers = tipo != "spec_script"
    cabecalho = str(dados.get("cabecalho") or "").strip()
    rodape = str(dados.get("rodape") or "").strip()
    header_footer = False
    if tipo in {"revised_draft", "production_draft", "continuity_script"}:
        header_footer = True
    elif tipo == "shooting_script":
        header_footer = bool(cabecalho or rodape)
    return {
        "tipo": tipo,
        "label": label,
        "scene_numbers": scene_numbers,
        "header_footer": header_footer,
        "cabecalho": cabecalho,
        "rodape": rodape,
    }


@registrar_etapa
def _texto_bloco_limpo(bloco: Tag) -> str:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        bloco: Bloco de conteúdo extraído ou gerado durante o processamento.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    return re.sub(r"\s+", " ", " ".join(bloco.stripped_strings)).strip()


@registrar_etapa
def _itens_unicos_normalizados(valores: list[str]) -> list[str]:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        valores: Valor usado pela rotina para compor a operação de itens unicos normalizados.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    vistos = set()
    itens = []
    for valor in valores:
        limpo = re.sub(r"\s+", " ", str(valor or "")).strip()
        if not limpo:
            continue
        chave = limpo.upper()
        if chave in vistos:
            continue
        vistos.add(chave)
        itens.append(limpo)
    return itens


@registrar_etapa
def _extrair_dados_cabecalho_cena(texto: str) -> dict | None:
    """
    Extrai informações de uma origem específica e devolve uma representação mais simples para as próximas etapas do fluxo.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        texto: Texto bruto ou parcialmente tratado que será analisado pela função.

    Retorno:
        Dados extraídos em uma estrutura simplificada, adequada para as próximas etapas.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    texto_limpo = _strip_scene_number_prefix_text(texto)
    texto_normalizado = re.sub(r"\s+", " ", str(texto_limpo or "")).strip().upper()
    match = re.match(r"^(INT\./EXT\.|INT\.|EXT\.|CAP\s+\d+)(?:\s+(.*))?$", texto_normalizado, flags=re.UNICODE)
    if not match:
        return None
    prefixo = match.group(1)
    resto = (match.group(2) or "").strip()
    if " - " in resto:
        local, periodo = resto.split(" - ", 1)
        return {"prefixo": prefixo, "local": local.strip(), "periodo": periodo.strip()}
    return {"prefixo": prefixo, "local": resto.strip(), "periodo": ""}


@registrar_etapa
def _linhas_contato_roteiro(dados: dict) -> list[str]:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        dados: Valor usado pela rotina para compor a operação de linhas contato roteiro.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    linhas = []
    autor = re.sub(r"\s+", " ", str(dados.get("autor") or "")).strip()
    contatos = str(dados.get("contatos") or "").strip()
    informacoes = str(dados.get("informacoes_adicionais") or "").strip()
    if autor:
        linhas.append(autor)
    if contatos:
        linhas.extend([linha.strip() for linha in contatos.splitlines() if linha.strip()])
    if informacoes:
        linhas.extend([linha.strip() for linha in informacoes.splitlines() if linha.strip()])
    return _itens_unicos_normalizados(linhas)


@registrar_etapa
def _ano_capa_roteiro(dados: dict) -> str:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        dados: Valor usado pela rotina para compor a operação de ano capa roteiro.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    for chave in ("ano", "ano_projeto", "ano_roteiro"):
        valor = str(dados.get(chave) or "").strip()
        if valor:
            return valor
    return str(datetime.now().year)


@registrar_etapa
def _texto_copyright_capa(dados: dict) -> str:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        dados: Valor usado pela rotina para compor a operação de texto copyright capa.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    if not bool(dados.get("copyright")):
        return ""
    autor = re.sub(r"\s+", " ", str(dados.get("autor") or "")).strip()
    if not autor:
        return ""
    return f"© {_ano_capa_roteiro(dados)} {autor}"


@registrar_etapa
def _rodape_capa_roteiro(dados: dict) -> list[str]:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        dados: Valor usado pela rotina para compor a operação de rodape capa roteiro.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    linhas = []
    copyright = _texto_copyright_capa(dados)
    if copyright:
        linhas.append(copyright)
    linhas.extend(_linhas_contato_roteiro(dados))
    return _itens_unicos_normalizados(linhas)


@registrar_etapa
def _nome_arquivo_unicode(texto: str, fallback: str = "roteiro") -> str:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        texto: Texto bruto ou parcialmente tratado que será analisado pela função.
        fallback: Valor usado pela rotina para compor a operação de nome arquivo unicode.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    texto = re.sub(r'[\\/:*?"<>|]+', '', str(texto or '')).strip()
    texto = re.sub(r'\s+', ' ', texto)
    texto = texto.strip(' .')
    return texto or fallback


@registrar_etapa
def _blocos_roteiro_exportacao(html_fragmento: str, prefixo_cena: str = "0", numeracao_inicial: int = 1, numerar_cenas: bool = False) -> list[dict]:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        html_fragmento: Valor usado pela rotina para compor a operação de blocos roteiro exportacao.
        prefixo_cena: Valor usado pela rotina para compor a operação de blocos roteiro exportacao.
        numeracao_inicial: Valor usado pela rotina para compor a operação de blocos roteiro exportacao.
        numerar_cenas: Valor usado pela rotina para compor a operação de blocos roteiro exportacao.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    soup = _fragmento_html(html_fragmento)
    blocos = []
    indice_cena = max(1, int(numeracao_inicial or 1))
    prefixo = str(prefixo_cena or "").strip()
    for bloco in _iterar_blocos(soup):
        tipo = _tipo_bloco_roteiro(bloco) or "action"
        if tipo == "comment":
            continue
        texto = _texto_bloco_limpo(bloco)
        if not texto:
            continue
        if tipo == "scene_heading":
            texto_base = _strip_scene_number_prefix_text(texto).upper()
            numero = None
            texto = texto_base
            if numerar_cenas:
                numero = f"{prefixo}{indice_cena}" if prefixo else str(indice_cena)
                texto = f"{numero}. {texto_base}"
                indice_cena += 1
        else:
            numero = None
            if tipo in {"character", "shot", "transition"}:
                texto = texto.upper()
        blocos.append({"tipo": tipo, "texto": texto, "numero_cena": numero})
    return blocos


@registrar_etapa
def _estimar_segundos_dialogo(falas: list[str]) -> float:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        falas: Valor usado pela rotina para compor a operação de estimar segundos dialogo.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    total_segundos = 0.0
    for texto in falas:
        palavras = len(re.findall(r"\b[\wÀ-ÿ'-]+\b", texto or "", flags=re.UNICODE))
        subtotal = 0.0
        subtotal += palavras / 3.0
        subtotal += 0.10 * len(re.findall(r"[,;:]", texto or ""))
        subtotal += 0.30 * len(re.findall(r"\.\.\.|…", texto or ""))
        sem_reticencias = re.sub(r"\.\.\.|…", "", texto or "")
        subtotal += 0.25 * len(re.findall(r"[.!?]", sem_reticencias))
        if subtotal > 0 and subtotal < 1:
            subtotal = 0.5
        else:
            subtotal = round(subtotal * 2) / 2
        total_segundos += subtotal
    return max(0.0, total_segundos)


@registrar_etapa
def _formatar_duracao_segundos(segundos: int | float) -> str:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        segundos: Valor usado pela rotina para compor a operação de formatar duracao segundos.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    total = max(0, int(round(float(segundos or 0))))
    if total < 60:
        return f"{total} seg" if total != 1 else "1 seg"
    minutos, segs = divmod(total, 60)
    if segs:
        return f"{minutos} min {segs:02d} seg"
    return f"{minutos} min" if minutos != 1 else "1 min"


@registrar_etapa
def _coletar_estatisticas_roteiro(dados: dict, blocos: list[dict]) -> dict:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        dados: Valor usado pela rotina para compor a operação de coletar estatisticas roteiro.
        blocos: Valor usado pela rotina para compor a operação de coletar estatisticas roteiro.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    contador = Counter(bloco["tipo"] for bloco in blocos)
    textos = [bloco["texto"] for bloco in blocos]
    texto_total = "\n".join(textos).strip()
    palavras_total = len(re.findall(r"\b[\wÀ-ÿ'-]+\b", texto_total, flags=re.UNICODE))
    caracteres_total = len(texto_total.replace("\n", ""))
    caracteres_sem_espacos = len(re.sub(r"\s+", "", texto_total, flags=re.UNICODE))
    falas = [bloco["texto"] for bloco in blocos if bloco["tipo"] == "dialogue"]
    acoes = [bloco["texto"] for bloco in blocos if bloco["tipo"] == "action"]
    personagens_dialogo = Counter()
    personagem_atual = None
    personagens_encontrados = []
    for bloco in blocos:
        if bloco["tipo"] == "character":
            personagem_atual = re.sub(r"[()]", "", bloco["texto"]).strip().upper()
            if personagem_atual:
                personagens_encontrados.append(personagem_atual)
        elif bloco["tipo"] == "dialogue" and personagem_atual:
            personagens_dialogo[personagem_atual] += len(re.findall(r"\b[\wÀ-ÿ'-]+\b", bloco["texto"], flags=re.UNICODE))
        elif bloco["tipo"] not in {"parenthetical", "dialogue"}:
            personagem_atual = None
    cenas = [bloco for bloco in blocos if bloco["tipo"] == "scene_heading"]
    cenas_outline = [bloco["texto"] for bloco in cenas]
    dialogo_palavras = sum(len(re.findall(r"\b[\wÀ-ÿ'-]+\b", texto, flags=re.UNICODE)) for texto in falas)
    acao_palavras = sum(len(re.findall(r"\b[\wÀ-ÿ'-]+\b", texto, flags=re.UNICODE)) for texto in acoes)
    duracao_dialogo_seg = _estimar_segundos_dialogo(falas)
    duracao_dialogo_min = (duracao_dialogo_seg / 60.0) if duracao_dialogo_seg else 0
    duracao_acao_min = acao_palavras / 220 if acao_palavras else 0
    duracao_total_min = max((palavras_total / 180) if palavras_total else 0, duracao_dialogo_min + duracao_acao_min)

    personagens_catalogo = _itens_unicos_normalizados([str(item.get("nome") or "").strip() for item in (dados.get("catalogo_personagens") or [])])
    personagens_unicos = _itens_unicos_normalizados(personagens_catalogo + personagens_encontrados + list(personagens_dialogo.keys()))

    locais_cenas = []
    locais_uso = Counter()
    for cena in cenas:
        info = _extrair_dados_cabecalho_cena(cena["texto"])
        local = (info or {}).get("local") or _strip_scene_number_prefix_text(cena["texto"])
        local = re.sub(r"\s+", " ", str(local or "")).strip().upper()
        if local:
            locais_cenas.append(local)
            locais_uso[local] += 1
    locais_unicos = _itens_unicos_normalizados(locais_cenas)

    media_palavras_cena = (palavras_total / len(cenas)) if cenas else 0
    media_palavras_fala = (dialogo_palavras / len(falas)) if falas else 0
    return {
        "palavras_total": palavras_total,
        "caracteres_total": caracteres_total,
        "caracteres_sem_espacos_total": caracteres_sem_espacos,
        "cenas_total": len(cenas),
        "falas_total": len(falas),
        "blocos_acao_total": len(acoes),
        "personagens_total": len(personagens_unicos),
        "locais_total": len(locais_unicos),
        "duracao_estimada_min": duracao_total_min,
        "duracao_estimada_texto": _formatar_duracao_estimada(duracao_total_min),
        "duracao_dialogo_segundos": duracao_dialogo_seg,
        "duracao_dialogo_texto": _formatar_duracao_segundos(duracao_dialogo_seg),
        "palavras_dialogo_total": dialogo_palavras,
        "palavras_acao_total": acao_palavras,
        "media_palavras_por_cena": media_palavras_cena,
        "media_palavras_por_fala": media_palavras_fala,
        "top_personagens": personagens_dialogo.most_common(8),
        "locais_uso": locais_uso.most_common(12),
        "outline": cenas_outline,
        "tipos_bloco": dict(contador),
    }


@registrar_etapa
def _formatar_duracao_estimada(minutos: float) -> str:
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        minutos: Valor usado pela rotina para compor a operação de formatar duracao estimada.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    total = max(0, int(round(float(minutos or 0))))
    if total < 60:
        return f"{total} min" if total != 1 else "1 min"
    horas, mins = divmod(total, 60)
    return f"{horas}h {mins:02d} min" if mins else f"{horas}h"


@registrar_etapa
def _aplicar_fonte_paragrafo_docx(paragrafo, fonte: str = "Courier New", tamanho: float = 12, negrito: bool = False, italico: bool = False):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        paragrafo: Valor usado pela rotina para compor a operação de aplicar fonte paragrafo docx.
        fonte: Valor usado pela rotina para compor a operação de aplicar fonte paragrafo docx.
        tamanho: Valor usado pela rotina para compor a operação de aplicar fonte paragrafo docx.
        negrito: Valor usado pela rotina para compor a operação de aplicar fonte paragrafo docx.
        italico: Valor usado pela rotina para compor a operação de aplicar fonte paragrafo docx.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    for run in paragrafo.runs:
        _run_set_font(run, fonte)
        run.font.size = Pt(tamanho)
        run.bold = negrito or bool(run.bold)
        run.italic = italico or bool(run.italic)
    if not paragrafo.runs:
        run = paragrafo.add_run("")
        _run_set_font(run, fonte)
        run.font.size = Pt(tamanho)
        run.bold = negrito
        run.italic = italico


@registrar_etapa
def _adicionar_heading_docx(doc: Document, texto: str, nivel: int = 1):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        doc: Documento em processamento, normalmente vindo de uma biblioteca de exportação ou leitura.
        texto: Texto bruto ou parcialmente tratado que será analisado pela função.
        nivel: Valor usado pela rotina para compor a operação de adicionar heading docx.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if nivel > 1 else WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12 if nivel > 1 else 18)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(texto)
    run.bold = True
    _run_set_font(run, "Courier New")
    run.font.size = Pt(18 if nivel == 1 else 14)
    return p


@registrar_etapa
def _adicionar_capa_roteiro_docx(doc: Document, projeto: dict, dados: dict, config: dict):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        doc: Documento em processamento, normalmente vindo de uma biblioteca de exportação ou leitura.
        projeto: Dicionário com os metadados e configurações do projeto atual.
        dados: Valor usado pela rotina para compor a operação de adicionar capa roteiro docx.
        config: Valor usado pela rotina para compor a operação de adicionar capa roteiro docx.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    titulo_texto = (dados.get("titulo_roteiro") or dados.get("titulo_projeto") or "SEM TÍTULO").upper()
    linhas_rodape = _rodape_capa_roteiro(dados)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.paragraph_format.space_before = Pt(180)
    titulo.paragraph_format.space_after = Pt(18)
    run = titulo.add_run(titulo_texto)
    run.bold = True
    run.font.size = Pt(12)
    _run_set_font(run, "Courier New")

    autor = str(dados.get("autor") or projeto.get("autor") or "").strip()
    if autor:
        escrito_por = doc.add_paragraph()
        escrito_por.alignment = WD_ALIGN_PARAGRAPH.CENTER
        escrito_por.paragraph_format.space_after = Pt(2)
        run = escrito_por.add_run("Escrito por")
        _run_set_font(run, "Courier New")
        run.font.size = Pt(12)

        autor_nome = doc.add_paragraph()
        autor_nome.alignment = WD_ALIGN_PARAGRAPH.CENTER
        autor_nome.paragraph_format.space_after = Pt(0)
        run = autor_nome.add_run(autor)
        _run_set_font(run, "Courier New")
        run.font.size = Pt(12)

    if linhas_rodape:
        espaco = doc.add_paragraph()
        espaco.paragraph_format.space_before = Pt(170)
        espaco.paragraph_format.space_after = Pt(0)
        for idx, linha in enumerate(linhas_rodape):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0 if idx == len(linhas_rodape) - 1 else 2)
            run = p.add_run(linha)
            _run_set_font(run, "Courier New")
            run.font.size = Pt(11 if idx == 0 and linha.startswith("©") else 12)




@registrar_etapa
def _adicionar_contra_capa_roteiro_docx(doc: Document, dados: dict):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        doc: Documento em processamento, normalmente vindo de uma biblioteca de exportação ou leitura.
        dados: Valor usado pela rotina para compor a operação de adicionar contra capa roteiro docx.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_before = Pt(170)
    p_titulo.paragraph_format.space_after = Pt(18)
    run = p_titulo.add_run("FICHA TÉCNICA")
    _run_set_font(run, "Courier New")
    run.font.size = Pt(12)
    run.bold = True

    linhas = []
    if dados.get("tipo_roteiro_label"):
        linhas.append(f"Tipo: {dados['tipo_roteiro_label']}")
    if dados.get("genero"):
        linhas.append(f"Gênero: {dados['genero']}")
    if dados.get("logline"):
        linhas.append(f"Logline: {dados['logline']}")
    if dados.get("sinopse"):
        linhas.append(f"Sinopse: {dados['sinopse']}")
    for linha in linhas:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(1.2)
        p.paragraph_format.right_indent = Inches(1.0)
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(linha)
        _run_set_font(run, "Courier New")
        run.font.size = Pt(12)
@registrar_etapa
def _adicionar_outline_docx(doc: Document, blocos: list[dict], config: dict):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        doc: Documento em processamento, normalmente vindo de uma biblioteca de exportação ou leitura.
        blocos: Valor usado pela rotina para compor a operação de adicionar outline docx.
        config: Valor usado pela rotina para compor a operação de adicionar outline docx.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    cenas = [bloco for bloco in blocos if bloco["tipo"] == "scene_heading"]
    if not cenas:
        p = doc.add_paragraph("Nenhuma cena identificada.")
        _aplicar_fonte_paragrafo_docx(p)
        return
    for idx, cena in enumerate(cenas, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        texto = cena["texto"]
        if not config["scene_numbers"]:
            texto = _strip_scene_number_prefix_text(texto)
        run = p.add_run(texto)
        _run_set_font(run, "Courier New")
        run.font.size = Pt(12)


@registrar_etapa
def _adicionar_estatisticas_docx(doc: Document, stats: dict, dados: dict | None = None):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        doc: Documento em processamento, normalmente vindo de uma biblioteca de exportação ou leitura.
        stats: Valor usado pela rotina para compor a operação de adicionar estatisticas docx.
        dados: Valor usado pela rotina para compor a operação de adicionar estatisticas docx.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    linhas = [
        ("Cenas", stats["cenas_total"]),
        ("Palavras", stats["palavras_total"]),
        ("Caracteres", stats["caracteres_total"]),
        ("Caracteres sem espaços", stats["caracteres_sem_espacos_total"]),
        ("Blocos de ação", stats["blocos_acao_total"]),
        ("Blocos de diálogo", stats["falas_total"]),
        ("Palavras em ação", stats["palavras_acao_total"]),
        ("Palavras em diálogo", stats["palavras_dialogo_total"]),
        ("Personagens", stats["personagens_total"]),
        ("Locais", stats["locais_total"]),
        ("Duração estimada", stats["duracao_estimada_texto"]),
        ("Tempo de diálogo", stats["duracao_dialogo_texto"]),
        ("Média de palavras por cena", f"{stats['media_palavras_por_cena']:.1f}"),
        ("Média de palavras por fala", f"{stats['media_palavras_por_fala']:.1f}"),
    ]
    for rotulo, valor in linhas:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{rotulo}: ")
        r1.bold = True
        _run_set_font(r1, "Courier New")
        r1.font.size = Pt(11)
        r2 = p.add_run(str(valor))
        _run_set_font(r2, "Courier New")
        r2.font.size = Pt(11)
    if stats.get("locais_uso"):
        for nome, cenas in stats["locais_uso"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{nome}: {cenas} cena(s)")
            _run_set_font(run, "Courier New")
            run.font.size = Pt(11)
    if stats.get("top_personagens"):
        for nome, palavras in stats["top_personagens"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{nome}: {palavras} palavras em diálogo")
            _run_set_font(run, "Courier New")
            run.font.size = Pt(11)


@registrar_etapa
def _limpar_secao_docx(sec, usar_header_footer: bool = False, cabecalho: str = "", rodape: str = ""):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        sec: Valor usado pela rotina para compor a operação de limpar secao docx.
        usar_header_footer: Valor usado pela rotina para compor a operação de limpar secao docx.
        cabecalho: Valor usado pela rotina para compor a operação de limpar secao docx.
        rodape: Valor usado pela rotina para compor a operação de limpar secao docx.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    if usar_header_footer:
        _configurar_secao_roteiro_docx(sec, cabecalho, rodape)
    else:
        _configurar_secao_roteiro_docx(sec, "", "")
        sec.header.paragraphs[0].clear()
        sec.first_page_header.paragraphs[0].clear()
        sec.footer.paragraphs[0].clear()
        sec.first_page_footer.paragraphs[0].clear()


@registrar_etapa
def _adicionar_bloco_roteiro_docx(doc: Document, bloco: dict):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        doc: Documento em processamento, normalmente vindo de uma biblioteca de exportação ou leitura.
        bloco: Bloco de conteúdo extraído ou gerado durante o processamento.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    tipo = bloco["tipo"]
    texto = bloco["texto"]
    recuo_esq, recuo_dir = ROTEIRO_RECUOS_POL.get(tipo, ROTEIRO_RECUOS_POL["action"])
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(ROTEIRO_ESPACOS_DEPOIS_PT.get(tipo, ROTEIRO_ESPACO_PADRAO_PT))
    fmt.line_spacing = 1.0
    fmt.first_line_indent = Pt(0)
    fmt.left_indent = Inches(recuo_esq)
    fmt.right_indent = Inches(recuo_dir)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if tipo == "transition" else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(texto)
    _run_set_font(run, ROTEIRO_FONTE_DOCX)
    run.font.size = Pt(ROTEIRO_TAMANHO_PT)
    run.bold = False


@registrar_etapa
def _adicionar_roteiro_docx(doc: Document, blocos: list[dict]):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        doc: Documento em processamento, normalmente vindo de uma biblioteca de exportação ou leitura.
        blocos: Valor usado pela rotina para compor a operação de adicionar roteiro docx.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    for bloco in blocos:
        _adicionar_bloco_roteiro_docx(doc, bloco)


@registrar_etapa
def exportar_roteiro_docx(slug: str, numero_roteiro: int | None = None, tipo_documento: str = "roteiro") -> Path:
    """
    Conduz a exportação dos dados do projeto para o formato esperado, preparando o conteúdo e delegando etapas auxiliares quando necessário.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        slug: Identificador amigável do projeto, livro ou recurso que será manipulado.
        numero_roteiro: Valor usado pela rotina para compor a operação de exportar roteiro docx.
        tipo_documento: Valor usado pela rotina para compor a operação de exportar roteiro docx.

    Retorno:
        O caminho do arquivo exportado ou uma estrutura com informações da exportação.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    projeto = obter_projeto(slug)
    pasta_projeto = obter_pasta_projeto(slug)
    pasta_destino = obter_pasta_exportacao_projeto(projeto["titulo"])

    roteiros = projeto.get("roteiros", [])
    if numero_roteiro is not None:
        roteiros = [item for item in roteiros if int(item.get("id", 0) or 0) == int(numero_roteiro)]
        if not roteiros:
            raise FileNotFoundError("Roteiro não encontrado para exportação.")
    if not roteiros:
        raise FileNotFoundError("Nenhum roteiro encontrado para exportação.")

    item = roteiros[0]
    roteiro = ler_roteiro(pasta_projeto, item["id"])
    dados = _coletar_dados_roteiro_profissional(projeto, roteiro)
    config = _tipo_exportacao_roteiro(dados)
    blocos = _blocos_roteiro_exportacao(dados["html"], dados["prefixo_cena"], dados["numeracao_inicial"], config["scene_numbers"])
    stats = _coletar_estatisticas_roteiro(dados, blocos)

    tipo_documento = (tipo_documento or "roteiro").strip().lower()
    opcoes_validas = {"roteiro", "outline", "estatistica", "estatística", "completo"}
    if tipo_documento not in opcoes_validas:
        raise ValueError("Tipo de documento inválido.")
    if tipo_documento == "estatística":
        tipo_documento = "estatistica"

    titulo_base = _nome_arquivo_unicode(dados.get("titulo_roteiro") or projeto.get("titulo") or "roteiro", "roteiro")

    doc = Document()
    sec = doc.sections[0]
    _limpar_secao_docx(sec, False)
    normal = doc.styles["Normal"]
    normal.font.name = ROTEIRO_FONTE_DOCX
    normal.font.size = Pt(ROTEIRO_TAMANHO_PT)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(0)

    _adicionar_capa_roteiro_docx(doc, projeto, dados, config)

    if tipo_documento == "completo":
        doc.add_page_break()
        _limpar_secao_docx(doc.sections[-1], False)
        _adicionar_outline_docx(doc, blocos, config)
        doc.add_section(WD_SECTION.NEW_PAGE)
        _limpar_secao_docx(doc.sections[-1], config["header_footer"], config["cabecalho"], config["rodape"])
        _adicionar_roteiro_docx(doc, blocos)
        doc.add_section(WD_SECTION.NEW_PAGE)
        _limpar_secao_docx(doc.sections[-1], False)
        _adicionar_estatisticas_docx(doc, stats, dados)
        caminho_saida = pasta_destino / f"{titulo_base}_completo.docx"
    elif tipo_documento == "outline":
        doc.add_page_break()
        _adicionar_outline_docx(doc, blocos, config)
        caminho_saida = pasta_destino / f"{titulo_base}_outline.docx"
    elif tipo_documento == "estatistica":
        doc.add_page_break()
        _adicionar_estatisticas_docx(doc, stats, dados)
        caminho_saida = pasta_destino / f"{titulo_base}_estatística.docx"
    else:
        doc.add_page_break()
        _limpar_secao_docx(doc.sections[0], config["header_footer"], config["cabecalho"], config["rodape"])
        _adicionar_roteiro_docx(doc, blocos)
        caminho_saida = pasta_destino / f"{titulo_base}_roteiro.docx"

    doc.save(caminho_saida)
    return caminho_saida


@registrar_etapa
def _pdf_estilo(base, nome: str, **kwargs):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        base: Valor usado pela rotina para compor a operação de pdf estilo.
        nome: Valor usado pela rotina para compor a operação de pdf estilo.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    return ParagraphStyle(name=nome, parent=base, **kwargs)


@registrar_etapa
def _adicionar_capa_roteiro_pdf(story: list, estilos: StyleSheet1, projeto: dict, dados: dict, config: dict):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        story: Valor usado pela rotina para compor a operação de adicionar capa roteiro pdf.
        estilos: Valor usado pela rotina para compor a operação de adicionar capa roteiro pdf.
        projeto: Dicionário com os metadados e configurações do projeto atual.
        dados: Valor usado pela rotina para compor a operação de adicionar capa roteiro pdf.
        config: Valor usado pela rotina para compor a operação de adicionar capa roteiro pdf.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    titulo_texto = (dados.get("titulo_roteiro") or dados.get("titulo_projeto") or "SEM TÍTULO").upper()
    autor = str(dados.get("autor") or projeto.get("autor") or "").strip()
    linhas_rodape = _rodape_capa_roteiro(dados)

    estilo_titulo = _pdf_estilo(estilos["Base"], f"TituloCapa_{len(story)}", fontName="Courier-Bold", fontSize=16, leading=20, alignment=TA_CENTER, firstLineIndent=0, spaceAfter=20)
    estilo_centro = _pdf_estilo(estilos["Base"], f"CentroCapa_{len(story)}", fontName="Courier", fontSize=12, leading=14, alignment=TA_CENTER, firstLineIndent=0, spaceAfter=4)
    estilo_rodape = _pdf_estilo(estilos["Base"], f"RodapeCapa_{len(story)}", fontName="Courier", fontSize=11, leading=13, alignment=TA_LEFT, firstLineIndent=0, spaceAfter=2)

    story.append(Spacer(1, 175))
    story.append(Paragraph(html.escape(titulo_texto), estilo_titulo))
    if autor:
        story.append(Spacer(1, 18))
        story.append(Paragraph("Escrito por", estilo_centro))
        story.append(Spacer(1, 6))
        story.append(Paragraph(html.escape(autor), estilo_centro))

    story.append(Spacer(1, 190 if linhas_rodape else 260))
    for idx, linha in enumerate(linhas_rodape):
        estilo_linha = estilo_rodape if not (idx == 0 and linha.startswith("©")) else _pdf_estilo(estilo_rodape, f"RodapeCapaCopy_{len(story)}_{idx}", fontSize=10.5, leading=12.5)
        story.append(Paragraph(html.escape(linha), estilo_linha))




@registrar_etapa
def _adicionar_contra_capa_roteiro_pdf(story: list, estilos: StyleSheet1, dados: dict):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        story: Valor usado pela rotina para compor a operação de adicionar contra capa roteiro pdf.
        estilos: Valor usado pela rotina para compor a operação de adicionar contra capa roteiro pdf.
        dados: Valor usado pela rotina para compor a operação de adicionar contra capa roteiro pdf.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    ano = _ano_capa_roteiro(dados)
    autor = re.sub(r"\s+", " ", str(dados.get("autor") or "")).strip()
    copyright = f"Todos os direitos reservados a {autor}" if autor else ""
    estilo = _pdf_estilo(estilos["Base"], f"ContraCapa_{len(story)}", fontName="Courier", fontSize=12, leading=15, alignment=TA_CENTER, firstLineIndent=0, spaceAfter=6)
    story.append(Spacer(1, 330))
    if ano:
        story.append(Paragraph(html.escape(ano), estilo))
    if copyright:
        story.append(Paragraph(html.escape(copyright), estilo))


@registrar_etapa
def _adicionar_outline_pdf(story: list, estilos: StyleSheet1, blocos: list[dict], config: dict):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        story: Valor usado pela rotina para compor a operação de adicionar outline pdf.
        estilos: Valor usado pela rotina para compor a operação de adicionar outline pdf.
        blocos: Valor usado pela rotina para compor a operação de adicionar outline pdf.
        config: Valor usado pela rotina para compor a operação de adicionar outline pdf.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    corpo = _pdf_estilo(estilos["Base"], f"OutlineCorpo_{len(story)}", fontName="Courier", fontSize=12, leading=15, alignment=TA_LEFT, firstLineIndent=0, leftIndent=0, spaceAfter=8)
    cenas = [bloco for bloco in blocos if bloco["tipo"] == "scene_heading"]
    if not cenas:
        story.append(Paragraph("Nenhuma cena identificada.", corpo))
        return
    for cena in cenas:
        texto = cena["texto"] if config["scene_numbers"] else _strip_scene_number_prefix_text(cena["texto"])
        story.append(Paragraph(html.escape(texto), corpo))


@registrar_etapa
def _adicionar_roteiro_pdf(story: list, estilos: StyleSheet1, blocos: list[dict]):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        story: Valor usado pela rotina para compor a operação de adicionar roteiro pdf.
        estilos: Valor usado pela rotina para compor a operação de adicionar roteiro pdf.
        blocos: Valor usado pela rotina para compor a operação de adicionar roteiro pdf.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    for bloco in blocos:
        tipo = bloco["tipo"]
        base = estilos["ScreenplayBase"]
        nome_estilo = {
            "scene_heading": "ScreenplayScene",
            "action": "ScreenplayAction",
            "character": "ScreenplayCharacter",
            "dialogue": "ScreenplayDialogue",
            "parenthetical": "ScreenplayParenthetical",
            "shot": "ScreenplayShot",
            "transition": "ScreenplayTransition",
            "music": "ScreenplayMusic",
        }.get(tipo, "ScreenplayAction")
        estilo = _pdf_estilo(estilos[nome_estilo], f"{nome_estilo}_{len(story)}")
        story.append(Paragraph(html.escape(bloco["texto"]) or "&#160;", estilo))


@registrar_etapa
def _adicionar_estatisticas_pdf(story: list, estilos: StyleSheet1, stats: dict, dados: dict | None = None):
    """
    Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        story: Valor usado pela rotina para compor a operação de adicionar estatisticas pdf.
        estilos: Valor usado pela rotina para compor a operação de adicionar estatisticas pdf.
        stats: Valor usado pela rotina para compor a operação de adicionar estatisticas pdf.
        dados: Valor usado pela rotina para compor a operação de adicionar estatisticas pdf.

    Retorno:
        O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    corpo = _pdf_estilo(estilos["Base"], f"StatsCorpo_{len(story)}", fontName="Courier", fontSize=11, leading=14, alignment=TA_LEFT, firstLineIndent=0, spaceAfter=4)
    if dados:
        if str(dados.get("logline") or "").strip():
            story.append(Paragraph(f"<b>Logline:</b> {html.escape(str(dados.get('logline') or '').strip())}", corpo))
            story.append(Spacer(1, 8))
        if str(dados.get("sinopse") or "").strip():
            story.append(Paragraph(f"<b>Sinopse:</b> {html.escape(str(dados.get('sinopse') or '').strip())}", corpo))
            story.append(Spacer(1, 10))
    linhas = [
        ("Cenas", stats["cenas_total"]),
        ("Palavras", stats["palavras_total"]),
        ("Caracteres", stats["caracteres_total"]),
        ("Caracteres sem espaços", stats["caracteres_sem_espacos_total"]),
        ("Blocos de ação", stats["blocos_acao_total"]),
        ("Blocos de diálogo", stats["falas_total"]),
        ("Palavras em ação", stats["palavras_acao_total"]),
        ("Palavras em diálogo", stats["palavras_dialogo_total"]),
        ("Personagens", stats["personagens_total"]),
        ("Locais", stats["locais_total"]),
        ("Duração estimada", stats["duracao_estimada_texto"]),
        ("Tempo de diálogo", stats["duracao_dialogo_texto"]),
        ("Média de palavras por cena", f"{stats['media_palavras_por_cena']:.1f}"),
        ("Média de palavras por fala", f"{stats['media_palavras_por_fala']:.1f}"),
    ]
    for rotulo, valor in linhas:
        story.append(Paragraph(f"<b>{html.escape(rotulo)}:</b> {html.escape(str(valor))}", corpo))
    if stats.get("locais_uso"):
        story.append(Spacer(1, 10))
        story.append(Paragraph("<b>Uso de locais</b>", corpo))
        for nome, cenas in stats["locais_uso"]:
            story.append(Paragraph(html.escape(f"{nome}: {cenas} cena(s)"), corpo))
    if stats.get("top_personagens"):
        story.append(Spacer(1, 10))
        story.append(Paragraph("<b>Pesquisa de personagens</b>", corpo))
        for nome, palavras in stats["top_personagens"]:
            story.append(Paragraph(html.escape(f"{nome}: {palavras} palavras em diálogo"), corpo))


@registrar_etapa
def exportar_roteiro_pdf(slug: str, numero_roteiro: int | None = None, tipo_documento: str = "roteiro") -> Path:
    """
    Conduz a exportação dos dados do projeto para o formato esperado, preparando o conteúdo e delegando etapas auxiliares quando necessário.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        slug: Identificador amigável do projeto, livro ou recurso que será manipulado.
        numero_roteiro: Valor usado pela rotina para compor a operação de exportar roteiro pdf.
        tipo_documento: Valor usado pela rotina para compor a operação de exportar roteiro pdf.

    Retorno:
        O caminho do arquivo exportado ou uma estrutura com informações da exportação.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    projeto = obter_projeto(slug)
    pasta_projeto = obter_pasta_projeto(slug)
    cfg = obter_configuracoes()
    pasta_destino = obter_pasta_exportacao_projeto(projeto["titulo"])

    roteiros = projeto.get("roteiros", [])
    if numero_roteiro is not None:
        roteiros = [item for item in roteiros if int(item.get("id", 0) or 0) == int(numero_roteiro)]
        if not roteiros:
            raise FileNotFoundError("Roteiro não encontrado para exportação.")
    if not roteiros:
        raise FileNotFoundError("Nenhum roteiro encontrado para exportação.")

    item = roteiros[0]
    roteiro = ler_roteiro(pasta_projeto, item["id"])
    dados = _coletar_dados_roteiro_profissional(projeto, roteiro)
    config = _tipo_exportacao_roteiro(dados)
    blocos = _blocos_roteiro_exportacao(dados["html"], dados["prefixo_cena"], dados["numeracao_inicial"], config["scene_numbers"])
    stats = _coletar_estatisticas_roteiro(dados, blocos)
    estilos = _stylesheet_pdf(cfg)

    tipo_documento = (tipo_documento or "roteiro").strip().lower()
    opcoes_validas = {"roteiro", "outline", "estatistica", "estatística", "completo"}
    if tipo_documento not in opcoes_validas:
        raise ValueError("Tipo de documento inválido.")
    if tipo_documento == "estatística":
        tipo_documento = "estatistica"

    titulo_base = _nome_arquivo_unicode(dados.get("titulo_roteiro") or projeto.get("titulo") or "roteiro", "roteiro")

    @registrar_etapa
    def _build_pdf(caminho: Path, montar_story, usar_callbacks: bool = False, ocultar_primeira_capa: bool = False, deslocamento_numero: int = 0):
        """
        Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.
    
        A função isola esta responsabilidade para deixar o fluxo principal mais fácil
        de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
        normalizações ou verificações necessárias e entrega um resultado previsível
        para a próxima camada do sistema.
    
        Parâmetros:
            caminho: Caminho de arquivo ou pasta usado como origem ou destino da operação.
            montar_story: Valor usado pela rotina para compor a operação de build pdf.
            usar_callbacks: Valor usado pela rotina para compor a operação de build pdf.
            ocultar_primeira_capa: Valor usado pela rotina para compor a operação de build pdf.
            deslocamento_numero: Valor usado pela rotina para compor a operação de build pdf.
    
        Retorno:
            O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.
    
        Observações:
            Mantém a lógica centralizada e evita que detalhes de implementação vazem
            para quem apenas precisa acionar este comportamento.
        """
        doc = SimpleDocTemplate(str(caminho), pagesize=(8.5 * inch, 11 * inch), leftMargin=1.5 * inch, rightMargin=1.0 * inch, topMargin=1.0 * inch, bottomMargin=1.0 * inch)
        story = []
        montar_story(story)
        if usar_callbacks and config["header_footer"]:
            primeira_cb, demais_cb = _criar_callbacks_roteiro_pdf(config["cabecalho"], config["rodape"], ocultar_primeira=ocultar_primeira_capa, deslocamento_numero=deslocamento_numero)
            if primeira_cb and demais_cb:
                doc.build(story, onFirstPage=primeira_cb, onLaterPages=demais_cb)
                return
        doc.build(story)

    if tipo_documento == "roteiro":
        caminho = pasta_destino / f"{titulo_base}_roteiro.pdf"
        _build_pdf(caminho, lambda story: (_adicionar_capa_roteiro_pdf(story, estilos, projeto, dados, config), story.append(PageBreak()), _adicionar_roteiro_pdf(story, estilos, blocos)), usar_callbacks=True, ocultar_primeira_capa=True, deslocamento_numero=1)
        return caminho
    if tipo_documento == "outline":
        caminho = pasta_destino / f"{titulo_base}_outline.pdf"
        _build_pdf(caminho, lambda story: (_adicionar_capa_roteiro_pdf(story, estilos, projeto, dados, config), story.append(PageBreak()), _adicionar_outline_pdf(story, estilos, blocos, config)))
        return caminho
    if tipo_documento == "estatistica":
        caminho = pasta_destino / f"{titulo_base}_estatística.pdf"
        _build_pdf(caminho, lambda story: (_adicionar_capa_roteiro_pdf(story, estilos, projeto, dados, config), story.append(PageBreak()), _adicionar_estatisticas_pdf(story, estilos, stats, dados)))
        return caminho

    caminho = pasta_destino / f"{titulo_base}_completo.pdf"
    _build_pdf(caminho, lambda story: (_adicionar_capa_roteiro_pdf(story, estilos, projeto, dados, config), story.append(PageBreak()), _adicionar_outline_pdf(story, estilos, blocos, config), story.append(PageBreak()), _adicionar_roteiro_pdf(story, estilos, blocos), story.append(PageBreak()), _adicionar_estatisticas_pdf(story, estilos, stats, dados)), usar_callbacks=True, ocultar_primeira_capa=True, deslocamento_numero=1)
    return caminho


@registrar_etapa
def exportar_estatisticas_epub_pdf(slug: str, stats: dict) -> Path:
    """
    Conduz a exportação dos dados do projeto para o formato esperado, preparando o conteúdo e delegando etapas auxiliares quando necessário.

    A função isola esta responsabilidade para deixar o fluxo principal mais fácil
    de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
    normalizações ou verificações necessárias e entrega um resultado previsível
    para a próxima camada do sistema.

    Parâmetros:
        slug: Identificador amigável do projeto, livro ou recurso que será manipulado.
        stats: Valor usado pela rotina para compor a operação de exportar estatisticas epub pdf.

    Retorno:
        O caminho do arquivo exportado ou uma estrutura com informações da exportação.

    Observações:
        Mantém a lógica centralizada e evita que detalhes de implementação vazem
        para quem apenas precisa acionar este comportamento.
    """
    projeto = obter_projeto(slug)
    pasta_destino = obter_pasta_exportacao_projeto(projeto["titulo"])
    caminho_saida = pasta_destino / f"{nome_seguro(projeto['titulo'])}_estatisticas.pdf"

    doc = SimpleDocTemplate(
        str(caminho_saida),
        pagesize=A4,
        leftMargin=1.4 * cm,
        rightMargin=1.4 * cm,
        topMargin=1.4 * cm,
        bottomMargin=1.4 * cm,
    )

    estilos = getSampleStyleSheet()
    titulo = ParagraphStyle(
        name='StatsTitulo',
        parent=estilos['Title'],
        fontName='Helvetica-Bold',
        fontSize=22,
        leading=26,
        textColor=colors.HexColor('#F8FAFC'),
        alignment=TA_LEFT,
        spaceAfter=8,
        wordWrap='CJK',
    )
    subtitulo = ParagraphStyle(
        name='StatsSubtitulo',
        parent=estilos['BodyText'],
        fontName='Helvetica',
        fontSize=11,
        leading=14,
        textColor=colors.HexColor('#D6E3FF'),
        spaceAfter=12,
        wordWrap='CJK',
    )
    secao = ParagraphStyle(
        name='StatsSecao',
        parent=estilos['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        textColor=colors.HexColor('#F8FAFC'),
        spaceBefore=12,
        spaceAfter=10,
        wordWrap='CJK',
    )
    texto = ParagraphStyle(
        name='StatsTexto',
        parent=estilos['BodyText'],
        fontName='Helvetica',
        fontSize=10,
        leading=13,
        textColor=colors.HexColor('#E2E8F0'),
        spaceAfter=5,
        wordWrap='CJK',
    )
    texto_mini = ParagraphStyle(
        name='StatsTextoMini',
        parent=texto,
        fontSize=9,
        leading=11,
        textColor=colors.HexColor('#CBD5E1'),
        wordWrap='CJK',
    )
    chip = ParagraphStyle(
        name='StatsChip',
        parent=texto,
        fontName='Helvetica-Bold',
        fontSize=9.5,
        leading=11,
        textColor=colors.HexColor('#F8FAFC'),
        alignment=TA_CENTER,
        spaceAfter=0,
        wordWrap='CJK',
    )

    resumo = stats.get('resumo') or {}
    ranking = stats.get('ranking_capitulos') or []
    top_vocab = stats.get('top_vocabulario') or []
    maior = stats.get('maior_capitulo')
    menor = stats.get('menor_capitulo')

    @registrar_etapa
    def p(txt, style=texto):
        """
        Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.
    
        A função isola esta responsabilidade para deixar o fluxo principal mais fácil
        de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
        normalizações ou verificações necessárias e entrega um resultado previsível
        para a próxima camada do sistema.
    
        Parâmetros:
            txt: Valor usado pela rotina para compor a operação de p.
            style: Valor usado pela rotina para compor a operação de p.
    
        Retorno:
            O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.
    
        Observações:
            Mantém a lógica centralizada e evita que detalhes de implementação vazem
            para quem apenas precisa acionar este comportamento.
        """
        return Paragraph(html.escape(str(txt or '—')).replace('\n', '<br/>'), style)

    @registrar_etapa
    def card(flowables, background='#111B45', padding=14, border='#334155'):
        """
        Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.
    
        A função isola esta responsabilidade para deixar o fluxo principal mais fácil
        de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
        normalizações ou verificações necessárias e entrega um resultado previsível
        para a próxima camada do sistema.
    
        Parâmetros:
            flowables: Valor usado pela rotina para compor a operação de card.
            background: Valor usado pela rotina para compor a operação de card.
            padding: Valor usado pela rotina para compor a operação de card.
            border: Valor usado pela rotina para compor a operação de card.
    
        Retorno:
            O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.
    
        Observações:
            Mantém a lógica centralizada e evita que detalhes de implementação vazem
            para quem apenas precisa acionar este comportamento.
        """
        bloco = Table([[flowables]], colWidths=[doc.width])
        bloco.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor(background)),
            ('BOX', (0, 0), (-1, -1), 0.8, colors.HexColor(border)),
            ('LEFTPADDING', (0, 0), (-1, -1), padding),
            ('RIGHTPADDING', (0, 0), (-1, -1), padding),
            ('TOPPADDING', (0, 0), (-1, -1), padding),
            ('BOTTOMPADDING', (0, 0), (-1, -1), padding),
        ]))
        return bloco

    @registrar_etapa
    def faixa_titulo(texto_secao: str):
        """
        Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.
    
        A função isola esta responsabilidade para deixar o fluxo principal mais fácil
        de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
        normalizações ou verificações necessárias e entrega um resultado previsível
        para a próxima camada do sistema.
    
        Parâmetros:
            texto_secao: Valor usado pela rotina para compor a operação de faixa titulo.
    
        Retorno:
            O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.
    
        Observações:
            Mantém a lógica centralizada e evita que detalhes de implementação vazem
            para quem apenas precisa acionar este comportamento.
        """
        barra = Table([[Paragraph(html.escape(texto_secao), secao)]], colWidths=[doc.width])
        barra.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#0D1433')),
            ('LINEBELOW', (0, 0), (-1, -1), 1.1, colors.HexColor('#6D5EF9')),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        return barra

    @registrar_etapa
    def desenhar_fundo(canvas, _doc):
        """
        Executa uma etapa específica do fluxo do módulo, encapsulando detalhes para manter o restante do código mais claro.
    
        A função isola esta responsabilidade para deixar o fluxo principal mais fácil
        de acompanhar. Ela recebe os dados no formato usado pelo projeto, realiza as
        normalizações ou verificações necessárias e entrega um resultado previsível
        para a próxima camada do sistema.
    
        Parâmetros:
            canvas: Valor usado pela rotina para compor a operação de desenhar fundo.
            _doc: Valor usado pela rotina para compor a operação de desenhar fundo.
    
        Retorno:
            O resultado produzido pela etapa, ou None quando a função atua apenas por efeito colateral.
    
        Observações:
            Mantém a lógica centralizada e evita que detalhes de implementação vazem
            para quem apenas precisa acionar este comportamento.
        """
        canvas.saveState()
        largura, altura = A4
        canvas.setFillColor(colors.HexColor('#050816'))
        canvas.rect(0, 0, largura, altura, fill=1, stroke=0)
        canvas.setFillColor(colors.HexColor('#0A1435'))
        canvas.circle(largura * 0.86, altura * 0.83, 110, fill=1, stroke=0)
        canvas.setFillColor(colors.HexColor('#101D4D'))
        canvas.circle(largura * 0.18, altura * 0.90, 85, fill=1, stroke=0)
        canvas.setFillColor(colors.HexColor('#0B1230'))
        canvas.rect(0, altura - 72, largura, 72, fill=1, stroke=0)
        canvas.restoreState()

    story = []

    hero_content = [
        Paragraph(html.escape(projeto.get('titulo') or 'Projeto'), titulo),
        Paragraph('Painel visual de estatísticas do EPUB', subtitulo),
        Paragraph(
            f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')} · leitura analítica do conteúdo e da estrutura editorial.",
            texto_mini,
        ),
    ]
    story.extend([card(hero_content, background='#111B45', padding=18), Spacer(1, 10)])

    kpis = [
        ('Palavras', f"{int(resumo.get('palavras') or 0):,}".replace(',', '.'), 'texto total'),
        ('Capítulos', f"{int(resumo.get('total_capitulos') or 0):,}".replace(',', '.'), 'estrutura ativa'),
        ('Leitura', str(resumo.get('tempo_leitura_legivel') or '0 min'), 'estimativa'),
        ('Vocabulário', f"{int(resumo.get('palavras_unicas') or 0):,}".replace(',', '.'), 'termos únicos'),
    ]
    kpi_rows = []
    current = []
    for label, valor, apoio in kpis:
        current.append([
            Paragraph(label.upper(), chip),
            Paragraph(html.escape(valor), ParagraphStyle(name=f'Valor{label}', parent=titulo, fontSize=18, leading=22, spaceAfter=4, wordWrap='CJK')),
            Paragraph(apoio, texto_mini),
        ])
    kpi_cells = []
    for label, valor, apoio in kpis:
        kpi_cells.append([
            Paragraph(label.upper(), chip),
            Paragraph(html.escape(valor), ParagraphStyle(name=f'Valor_{label}', parent=titulo, fontSize=18, leading=22, spaceAfter=4, wordWrap='CJK')),
            Paragraph(apoio, texto_mini),
        ])
    kpi_table_data = []
    row = []
    for cell in kpi_cells:
        row.append(cell)
    kpi_table_data.append(row)
    grade_kpi = Table(kpi_table_data, colWidths=[doc.width / 4.0] * 4)
    grade_kpi.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#192554')),
        ('BOX', (0, 0), (-1, -1), 0.8, colors.HexColor('#334155')),
        ('INNERGRID', (0, 0), (-1, -1), 0.8, colors.HexColor('#334155')),
        ('LEFTPADDING', (0, 0), (-1, -1), 12),
        ('RIGHTPADDING', (0, 0), (-1, -1), 12),
        ('TOPPADDING', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    story.extend([grade_kpi, Spacer(1, 12)])

    story.append(faixa_titulo('Panorama editorial'))
    panorama_rows = [['Métrica', 'Valor', 'Métrica', 'Valor']]
    panorama_rows += [
        [p('Parágrafos', texto_mini), p(resumo.get('paragrafos') or 0), p('Sentenças', texto_mini), p(resumo.get('sentencas') or 0)],
        [p('Média/capítulo', texto_mini), p(resumo.get('media_palavras_capitulo') or 0), p('Média/parágrafo', texto_mini), p(resumo.get('media_palavras_paragrafo') or 0)],
        [p('Imagens', texto_mini), p(resumo.get('imagens') or 0), p('Links', texto_mini), p(resumo.get('links') or 0)],
        [p('Listas', texto_mini), p(resumo.get('listas') or 0), p('Citações', texto_mini), p(resumo.get('citacoes') or 0)],
    ]
    tabela = Table(panorama_rows, colWidths=[doc.width * 0.34, doc.width * 0.16, doc.width * 0.34, doc.width * 0.16], repeatRows=1)
    tabela.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E2D63')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#0F172A')),
        ('GRID', (0, 0), (-1, -1), 0.6, colors.HexColor('#334155')),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    story.extend([tabela, Spacer(1, 12)])

    story.append(faixa_titulo('Destaques de capítulos'))
    destaque_cards = []
    if maior:
        destaque_cards.append([
            Paragraph('Mais extenso', chip),
            p(maior.get('titulo') or '—'),
            Paragraph(f"{int(maior.get('palavras') or 0):,} palavras · {maior.get('tempo_leitura_min') or 0} min".replace(',', '.'), texto_mini),
        ])
    if menor:
        destaque_cards.append([
            Paragraph('Mais enxuto', chip),
            p(menor.get('titulo') or '—'),
            Paragraph(f"{int(menor.get('palavras') or 0):,} palavras · {menor.get('tempo_leitura_min') or 0} min".replace(',', '.'), texto_mini),
        ])
    if not destaque_cards:
        destaque_cards.append([
            Paragraph('Capítulos', chip),
            p('Ainda não há conteúdo suficiente para análise detalhada.'),
            Paragraph('Adicione mais conteúdo ao projeto para enriquecer o relatório.', texto_mini),
        ])
    while len(destaque_cards) < 2:
        destaque_cards.append([Paragraph('', texto), Paragraph('', texto), Paragraph('', texto)])
    destaques_table = Table([destaque_cards], colWidths=[doc.width / 2.0] * 2)
    destaques_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#111827')),
        ('BOX', (0, 0), (-1, -1), 0.8, colors.HexColor('#334155')),
        ('INNERGRID', (0, 0), (-1, -1), 0.8, colors.HexColor('#334155')),
        ('LEFTPADDING', (0, 0), (-1, -1), 12),
        ('RIGHTPADDING', (0, 0), (-1, -1), 12),
        ('TOPPADDING', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    story.extend([destaques_table, Spacer(1, 12)])

    if ranking:
        story.append(faixa_titulo('Ranking de capítulos'))
        ranking_rows = [[
            Paragraph('Capítulo', chip),
            Paragraph('Palavras', chip),
            Paragraph('Participação', chip),
            Paragraph('Leitura', chip),
        ]]
        for item in ranking[:10]:
            ranking_rows.append([
                p(item.get('titulo') or f"Capítulo {item.get('id')}", texto),
                p(f"{int(item.get('palavras') or 0):,}".replace(',', '.'), texto),
                p(f"{item.get('participacao') or 0}%", texto),
                p(f"{item.get('tempo_leitura_min') or 0} min", texto),
            ])
        ranking_table = Table(ranking_rows, colWidths=[doc.width * 0.5, doc.width * 0.16, doc.width * 0.16, doc.width * 0.18], repeatRows=1)
        ranking_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E2D63')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#0F172A')),
            ('GRID', (0, 0), (-1, -1), 0.6, colors.HexColor('#334155')),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        story.extend([ranking_table, Spacer(1, 12)])

    if top_vocab:
        story.append(faixa_titulo('Vocabulário recorrente'))
        chip_rows = []
        row = []
        for palavra, qtd in top_vocab[:12]:
            row.append(Paragraph(f"{html.escape(str(palavra))} · {int(qtd)}", chip))
            if len(row) == 3:
                chip_rows.append(row)
                row = []
        if row:
            while len(row) < 3:
                row.append(Paragraph('', chip))
            chip_rows.append(row)
        chips_table = Table(chip_rows, colWidths=[doc.width / 3.0] * 3)
        chips_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#192554')),
            ('BOX', (0, 0), (-1, -1), 0.8, colors.HexColor('#334155')),
            ('INNERGRID', (0, 0), (-1, -1), 0.8, colors.HexColor('#334155')),
            ('LEFTPADDING', (0, 0), (-1, -1), 8),
            ('RIGHTPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        story.append(chips_table)

    doc.build(story, onFirstPage=desenhar_fundo, onLaterPages=desenhar_fundo)
    return caminho_saida
