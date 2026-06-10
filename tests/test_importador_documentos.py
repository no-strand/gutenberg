from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from modules import importador_documentos, manipulador_capitulos, manipulador_projetos, manipulador_roteiros
from tests.test_support import IsolatedProjectTestCase


class ImportadorDocumentosCapituloTests(IsolatedProjectTestCase):
    """Testes do importador de DOCX para capítulos de livro."""

    def _criar_docx_livro(self):
        caminho = self.base / "livro_importado.docx"
        doc = Document()

        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = titulo.add_run("Capítulo 1")
        run_titulo.bold = True
        run_titulo.font.size = None

        doc.add_paragraph("Primeiro parágrafo do capítulo.")
        doc.add_paragraph("Segundo parágrafo do capítulo, que precisa continuar separado.")
        doc.save(caminho)
        return caminho

    def test_docx_livro_importado_transforma_titulo_em_h1_e_preserva_paragrafos(self):
        """DOCX de livro: título vira h1 e cada parágrafo continua em seu próprio p."""
        manipulador_projetos.criar_projeto("Livro Importado", tipo="livro")
        caminho = self._criar_docx_livro()

        resultado = importador_documentos.importar_documento_como_capitulo("livro_importado", caminho)

        self.assertTrue(resultado["ok"])
        pasta = manipulador_projetos.obter_pasta_projeto("livro_importado")
        capitulo = manipulador_capitulos.ler_capitulo(pasta, 1)
        html = "".join(capitulo["paragrafos"])
        self.assertIn("<h1>Capítulo 1</h1>", html)
        self.assertIn("<p>Primeiro parágrafo do capítulo.</p>", html)
        self.assertIn("<p>Segundo parágrafo do capítulo, que precisa continuar separado.</p>", html)
        self.assertEqual(html.count("<p>"), 2)

    def test_importar_capitulo_recusa_projeto_de_roteiro(self):
        """Importação como capítulo: projeto de roteiro recebe erro claro."""
        manipulador_projetos.criar_projeto("Projeto Roteiro", tipo="roteiro")
        caminho = self._criar_docx_livro()

        with self.assertRaisesRegex(ValueError, "capítulo"):
            importador_documentos.importar_documento_como_capitulo("projeto_roteiro", caminho)


class ImportadorDocumentosRoteiroTests(IsolatedProjectTestCase):
    """Testes do importador de DOCX para roteiros."""

    def _criar_docx_roteiro(self):
        caminho = self.base / "roteiro_importado.docx"
        doc = Document()
        doc.add_paragraph("MEU ROTEIRO")
        doc.add_paragraph("Autor: Ana Silva")
        doc.add_paragraph("Contato: ana@example.com")
        doc.add_paragraph("INT. CASA - DIA")

        personagem = doc.add_paragraph("MARIA")
        personagem.alignment = WD_ALIGN_PARAGRAPH.CENTER
        personagem.runs[0].bold = True

        doc.add_paragraph("Olá, mundo.")
        doc.save(caminho)
        return caminho

    def test_docx_roteiro_importado_detecta_cabecalho_personagem_e_metadados(self):
        """DOCX de roteiro: importa blocos, autor/contato e catálogo detectável."""
        manipulador_projetos.criar_projeto("Roteiro Importado", tipo="roteiro")
        caminho = self._criar_docx_roteiro()

        resultado = importador_documentos.importar_documento_como_roteiro("roteiro_importado", caminho)

        self.assertTrue(resultado["ok"])
        projeto = manipulador_projetos.obter_projeto("roteiro_importado")
        self.assertEqual(projeto["autor"], "Ana Silva")
        self.assertEqual(projeto["contatos"], "ana@example.com")

        pasta = manipulador_projetos.obter_pasta_projeto("roteiro_importado")
        roteiro = manipulador_roteiros.ler_roteiro(pasta, 1)
        self.assertIn('data-block-type="scene_heading"', roteiro["html"])
        self.assertIn('data-block-type="character"', roteiro["html"])
        self.assertIn("MARIA", roteiro["personagens"])
        self.assertIn("CASA", roteiro["locais"])

    def test_importar_roteiro_recusa_projeto_de_livro(self):
        """Importação como roteiro: projeto de livro recebe erro claro."""
        manipulador_projetos.criar_projeto("Projeto Livro", tipo="livro")
        caminho = self._criar_docx_roteiro()

        with self.assertRaisesRegex(ValueError, "roteiro"):
            importador_documentos.importar_documento_como_roteiro("projeto_livro", caminho)
